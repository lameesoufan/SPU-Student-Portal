"""
Grades App — Service Layer

كل منطق العمل (Business Logic) الخاص بتطبيق العلامات موجود هون، بنفس نمط
تطبيقي accounts/projects. الـ Views لازم تبقى "غبية" (thin): تستقبل الطلب،
تتحقق من الصلاحيات الأساسية، تستدعي الدالة المناسبة من هون، وترجع الاستجابة.

كل دالة هون بترجع dict فيه على الأقل مفتاح 'ok' (True/False)، وبحالة الخطأ
'error' و 'status' (HTTP status code المقترح)، عشان الـ View يقدر يبني
Response مباشرة بدون ما يعرف تفاصيل منطق العمل.
"""
from __future__ import annotations

import io
from collections import defaultdict
from math import ceil

from django.contrib.auth import get_user_model
from django.db import transaction
from django.db.models import Q
from django.utils import timezone
from rest_framework import status

from committees.models import (
    Committee, ALL_COMMITTEE_TYPES, COMMITTEE_TYPE_AR, DEPARTMENT_AR, PROJECT_TYPE_AR,
)
from .models import (
    ProjectGrade, ProjectReport, GradeAuditLog, COMMITTEE_MAX_SCORES,
    CommitteeGradingMode, DoctorGradeDraft,
)
from .serializers import ProjectGradeSerializer, ProjectReportSerializer

User = get_user_model()


# ── Role / Ownership helpers ──────────────────────────────────────────────────
# (نفس الدوال يلي كانت بالـ views.py، منقولة بدون أي تغيير بالمنطق)

def is_student(user): return getattr(user, 'role', None) == 'student'
def is_doctor(user):  return getattr(user, 'role', None) in ('doctor', 'dean', 'hod')
def is_dean(user):    return getattr(user, 'role', None) == 'dean'
def is_hod(user):     return getattr(user, 'role', None) in ('hod', 'dean')


VALID_GRADE_PROJECT_SOURCES = {'IdeaApplication', 'StudentIdeaProposal'}


def committee_grader_ids(committee):
    """Return the current committee chair/member user IDs, without duplicates."""
    ids = set(committee.members.values_list('id', flat=True))
    if committee.chair_id:
        ids.add(committee.chair_id)
    return ids


def user_is_committee_grader(user, committee):
    """Only actual committee graders may read or submit individual drafts."""
    return bool(getattr(user, 'is_authenticated', False) and user.id in committee_grader_ids(committee))


def committee_contains_project(committee, source, pid):
    if source == 'IdeaApplication':
        return committee.applications.filter(pk=pid).exists()
    if source == 'StudentIdeaProposal':
        return committee.proposals.filter(pk=pid).exists()
    return False


def active_project_student_ids(source, pid):
    """Return active team members, with a legacy fallback only for projects
    that have no participation rows at all.

    Once ProjectParticipation rows exist, they are the source of truth even if
    every participant is failed or withdrawn. Falling back in that case would
    accidentally reactivate legacy leaders/invitees for grading.
    """
    from projects.models import IdeaApplication, ProjectParticipation, StudentIdeaProposal

    if source == 'IdeaApplication':
        project_participations = ProjectParticipation.objects.filter(idea_application_id=pid)
        ids = set(
            project_participations.filter(status='active')
            .values_list('student_id', flat=True)
        )
        if project_participations.exists():
            return ids
        project = (
            IdeaApplication.objects.filter(pk=pid)
            .select_related('student')
            .prefetch_related('invitations')
            .first()
        )
        if not project:
            return set()
        ids = {project.student_id} if project.student_id else set()
        ids.update(
            project.invitations.filter(status='accepted')
            .values_list('invitee_id', flat=True)
        )
        return ids

    if source == 'StudentIdeaProposal':
        project_participations = ProjectParticipation.objects.filter(student_proposal_id=pid)
        ids = set(
            project_participations.filter(status='active')
            .values_list('student_id', flat=True)
        )
        if project_participations.exists():
            return ids
        project = (
            StudentIdeaProposal.objects.filter(pk=pid)
            .select_related('student')
            .prefetch_related('invitations')
            .first()
        )
        if not project:
            return set()
        ids = {project.student_id} if project.student_id else set()
        ids.update(
            project.invitations.filter(status='accepted')
            .values_list('invitee_id', flat=True)
        )
        return ids

    return set()


def _normalise_grade_request(committee, source, pid, ctype, semester):
    if source not in VALID_GRADE_PROJECT_SOURCES:
        return None, 'مصدر المشروع غير صالح.'
    try:
        pid = int(pid)
    except (TypeError, ValueError):
        return None, 'project_id يجب أن يكون رقماً.'
    if ctype != committee.committee_type:
        return None, 'نوع اللجنة المرسل لا يطابق نوع اللجنة المحددة.'
    if semester and committee.semester and semester != committee.semester:
        return None, 'الفصل الدراسي المرسل لا يطابق فصل اللجنة.'
    if not committee_contains_project(committee, source, pid):
        return None, 'المشروع لا يتبع اللجنة المحددة.'
    return (pid, committee.semester or semester or ''), None


def get_project(source, pid):
    """جلب كائن المشروع بغض النظر عن المصدر."""
    from projects.models import IdeaApplication, StudentIdeaProposal
    if source == 'IdeaApplication':
        return IdeaApplication.objects.filter(pk=pid).first()
    if source == 'StudentIdeaProposal':
        return StudentIdeaProposal.objects.filter(pk=pid).first()
    return None


def student_belongs_to_project(user, source, pid):
    """هل الطالب عضو أو قائد في المشروع؟"""
    from projects.models import ProjectParticipation
    qs = ProjectParticipation.objects.filter(student=user, status='active')
    if source == 'IdeaApplication':
        return qs.filter(idea_application_id=pid).exists()
    return qs.filter(student_proposal_id=pid).exists()


def doctor_is_chair_for(user, source, pid, committee_type):
    """هل الدكتور رئيس لجنة من هذا النوع للمشروع؟"""
    qs = Committee.objects.filter(committee_type=committee_type, chair=user)
    if source == 'IdeaApplication':
        return qs.filter(applications__id=pid).exists()
    return qs.filter(proposals__id=pid).exists()


def doctor_is_member_for(user, source, pid, committee_type):
    """هل الدكتور/رئيس القسم عضو في لجنة من هذا النوع للمشروع؟"""
    qs = Committee.objects.filter(committee_type=committee_type)
    qs = qs.filter(Q(chair=user) | Q(members=user))
    if source == 'IdeaApplication':
        return qs.filter(applications__id=pid).exists()
    return qs.filter(proposals__id=pid).exists()


def resolve_grade_committee(source, pid, committee_type, semester='', committee_id=None):
    """Resolve the exact committee responsible for a grade-entry request.

    Direct grade entry must honor the grading mode of the *actual* committee.
    Therefore callers are not allowed to bypass mode checks simply by omitting
    ``committee_id``.  When the id is absent we resolve the committee from the
    project/type (and semester when supplied); ambiguous matches require an
    explicit id.
    """
    if source not in VALID_GRADE_PROJECT_SOURCES:
        return None, 'مصدر المشروع غير صالح.'

    try:
        pid = int(pid)
    except (TypeError, ValueError):
        return None, 'project_id يجب أن يكون رقماً.'

    if committee_id:
        try:
            committee = Committee.objects.prefetch_related('members').get(pk=committee_id)
        except (Committee.DoesNotExist, TypeError, ValueError):
            return None, 'اللجنة المحددة غير موجودة.'
        _, error = _normalise_grade_request(
            committee, source, pid, committee_type, semester
        )
        if error:
            return None, error
        return committee, None

    qs = Committee.objects.prefetch_related('members').filter(
        committee_type=committee_type,
    )
    if source == 'IdeaApplication':
        qs = qs.filter(applications__id=pid)
    else:
        qs = qs.filter(proposals__id=pid)
    if semester:
        qs = qs.filter(semester=semester)

    matches = list(qs.distinct()[:2])
    if not matches:
        return None, 'لا توجد لجنة مطابقة للمشروع ونوع التقييم المحددين.'
    if len(matches) > 1:
        return None, 'يوجد أكثر من لجنة مطابقة؛ يجب إرسال committee_id لتحديد اللجنة.'
    return matches[0], None


def project_q_filter(source, pid):
    if source == 'IdeaApplication':
        return Q(applications__id=pid)
    return Q(proposals__id=pid)


def doctor_can_access_report(user, source, pid):
    """هل الدكتور رئيس أو عضو في اللجنة النهائية لهذا المشروع؟"""
    qs = Committee.objects.filter(committee_type='final_discussion')
    qs = qs.filter(project_q_filter(source, pid))
    return qs.filter(Q(chair=user) | Q(members=user)).exists()


# ── Report Upload / Retrieval ─────────────────────────────────────────────────

ALLOWED_REPORT_EXTENSIONS = {'.pdf', '.doc', '.docx', '.zip', '.rar'}
ALLOWED_REPORT_MIME_TYPES = {
    '.pdf': {'application/pdf'},
    '.doc': {'application/msword'},
    '.docx': {'application/vnd.openxmlformats-officedocument.wordprocessingml.document'},
    '.zip': {'application/zip', 'application/x-zip-compressed'},
    '.rar': {'application/vnd.rar', 'application/x-rar-compressed', 'application/x-rar'},
}
MAX_REPORT_FILE_SIZE = 10 * 1024 * 1024  # 10 MB


def safe_uploaded_filename(filename):
    raw = str(filename or '').replace('\\', '/')
    safe = raw.rsplit('/', 1)[-1].strip()
    return safe[:255] or 'report'


def upload_report(*, user, source, pid, semester, file):
    """رفع تقرير مشروع من قبل الطالب صاحب المشروع فقط."""
    import os as _os

    if not is_student(user):
        return {'ok': False, 'error': 'مسموح للطلاب فقط.', 'status': status.HTTP_403_FORBIDDEN}

    if not (source and pid and file):
        return {'ok': False, 'error': 'project_source, project_id, file مطلوبة.',
                'status': status.HTTP_400_BAD_REQUEST}
    if source not in VALID_GRADE_PROJECT_SOURCES:
        return {'ok': False, 'error': 'مصدر المشروع غير صالح.',
                'status': status.HTTP_400_BAD_REQUEST}

    try:
        pid = int(pid)
    except (TypeError, ValueError):
        return {'ok': False, 'error': 'project_id يجب أن يكون رقماً.', 'status': status.HTTP_400_BAD_REQUEST}

    if not student_belongs_to_project(user, source, pid):
        return {'ok': False, 'error': 'لا تملك صلاحية رفع تقرير لهذا المشروع.',
                'status': status.HTTP_403_FORBIDDEN}

    if file.size > MAX_REPORT_FILE_SIZE:
        return {'ok': False, 'error': 'حجم الملف يتجاوز 10 MB.', 'status': status.HTTP_400_BAD_REQUEST}

    safe_name = safe_uploaded_filename(file.name)
    ext = _os.path.splitext(safe_name)[1].lower()
    if ext not in ALLOWED_REPORT_EXTENSIONS:
        return {
            'ok': False,
            'error': f'نوع الملف غير مسموح. المسموح: {", ".join(ALLOWED_REPORT_EXTENSIONS)}',
            'status': status.HTTP_400_BAD_REQUEST,
        }
    content_type = (getattr(file, 'content_type', '') or '').lower().split(';', 1)[0].strip()
    if content_type not in ALLOWED_REPORT_MIME_TYPES[ext]:
        return {'ok': False, 'error': 'نوع محتوى الملف لا يطابق امتداده.',
                'status': status.HTTP_400_BAD_REQUEST}
    file.name = safe_name

    report, created = ProjectReport.objects.get_or_create(
        project_source=source,
        project_id=pid,
        defaults={'semester': semester, 'uploaded_by': user},
    )
    if not created and report.file:
        try:
            report.file.delete(save=False)
        except Exception:
            pass

    report.file          = file
    report.original_name = safe_name
    report.file_size     = file.size
    report.semester       = semester or report.semester
    report.uploaded_by   = user
    report.save()

    return {
        'ok': True,
        'created': created,
        'report': report,
    }


def get_report_with_access_check(*, user, source, pid):
    """
    يرجّع التقرير إذا كان للمستخدم صلاحية الوصول له (طالب صاحب المشروع، أو
    دكتور عضو باللجنة النهائية، أو عميد)، وإلا يرجّع خطأ صلاحية.
    """
    pid = int(pid)

    if is_student(user):
        if not student_belongs_to_project(user, source, pid):
            return {'ok': False, 'error': 'ليس لديك صلاحية.', 'status': status.HTTP_403_FORBIDDEN}
    elif is_doctor(user):
        if not (doctor_can_access_report(user, source, pid) or is_dean(user)):
            return {'ok': False, 'error': 'ليس لديك صلاحية.', 'status': status.HTTP_403_FORBIDDEN}
    else:
        return {'ok': False, 'error': 'ليس لديك صلاحية.', 'status': status.HTTP_403_FORBIDDEN}

    try:
        report = ProjectReport.objects.get(project_source=source, project_id=pid)
    except ProjectReport.DoesNotExist:
        return {'ok': False, 'error': 'لم يُرفع التقرير بعد.', 'status': status.HTTP_404_NOT_FOUND}

    return {'ok': True, 'report': report}


# ── Enter Grade (single) ──────────────────────────────────────────────────────

def _check_grader_permission(user, source, pid, ctype):
    """Direct/individual grading is restricted to the committee chair only."""
    return doctor_is_chair_for(user, source, pid, ctype)


def enter_grade(*, user, validated_data):
    if not is_doctor(user):
        return {'ok': False, 'error': 'مسموح للدكاترة فقط.', 'status': status.HTTP_403_FORBIDDEN}

    d = validated_data
    source       = d['project_source']
    pid          = d['project_id']
    ctype        = d['committee_type']
    student_id   = d['student_id']
    committee_id = d.get('committee_id')
    semester     = d.get('semester', '')

    committee = None
    if committee_id:
        committee, error = resolve_grade_committee(
            source, pid, ctype, semester=semester, committee_id=committee_id
        )
        if error:
            return {'ok': False, 'error': error, 'status': status.HTTP_400_BAD_REQUEST}

    if committee is not None:
        allowed_individual_grader = committee.chair_id == user.id
    else:
        allowed_individual_grader = _check_grader_permission(user, source, pid, ctype)

    if not allowed_individual_grader:
        return {'ok': False, 'error': 'أنت لست رئيس اللجنة المسؤولة عن هذا المشروع.',
                'status': status.HTTP_403_FORBIDDEN}

    if committee is None:
        committee, error = resolve_grade_committee(
            source, pid, ctype, semester=semester, committee_id=None
        )
        if error:
            return {'ok': False, 'error': error, 'status': status.HTTP_400_BAD_REQUEST}

    mode = CommitteeGradingMode.objects.filter(committee=committee).first()
    if mode and mode.collective:
        return {
            'ok': False,
            'error': 'التقييم الجماعي مُفعَّل لهذه اللجنة؛ يجب إدخال تقييم كل عضو عبر مسار التقييم الجماعي.',
            'status': status.HTTP_409_CONFLICT,
        }

    committee_id = committee.id
    semester = committee.semester or semester

    try:
        student = User.objects.get(pk=student_id, role='student')
    except User.DoesNotExist:
        return {'ok': False, 'error': 'الطالب غير موجود.', 'status': status.HTTP_404_NOT_FOUND}
    if student.id not in active_project_student_ids(source, pid):
        return {'ok': False, 'error': 'الطالب ليس عضواً نشطاً في المشروع المحدد.',
                'status': status.HTTP_400_BAD_REQUEST}

    grade, created = ProjectGrade.objects.get_or_create(
        project_source=source, project_id=pid, committee_type=ctype, student=student,
        defaults={'semester': semester, 'committee_id': committee_id, 'entered_by': user},
    )

    if not created and not d.get('confirm_update', False):
        return {
            'ok': False,
            'requires_confirmation': True,
            'existing_grade': grade,
            'status': status.HTTP_409_CONFLICT,
        }

    old_main, old_report = grade.score_main, grade.score_report

    grade.score_main   = d['score_main']
    grade.score_report = d.get('score_report')
    grade.notes        = d.get('notes', '')
    grade.entered_by   = user
    if not grade.semester:
        grade.semester = semester
    grade.committee_id = committee_id
    grade.save()

    _log_grade_changes(grade, user, ctype, old_main, old_report)

    return {'ok': True, 'created': created, 'grade': grade}


def _log_grade_changes(grade, user, ctype, old_main, old_report):
    if old_main != grade.score_main:
        GradeAuditLog.objects.create(
            grade=grade, changed_by=user, field_changed='score_main',
            old_value=str(old_main) if old_main is not None else None,
            new_value=str(grade.score_main),
        )
    if ctype == 'final_discussion' and old_report != grade.score_report:
        GradeAuditLog.objects.create(
            grade=grade, changed_by=user, field_changed='score_report',
            old_value=str(old_report) if old_report is not None else None,
            new_value=str(grade.score_report),
        )


# ── Enter Bulk Grades ──────────────────────────────────────────────────────────

def enter_bulk_grades(*, user, validated_data):
    if not is_doctor(user):
        return {'ok': False, 'error': 'مسموح للدكاترة فقط.', 'status': status.HTTP_403_FORBIDDEN}

    d = validated_data
    source         = d['project_source']
    pid            = d['project_id']
    ctype          = d['committee_type']
    committee_id   = d.get('committee_id')
    semester       = d.get('semester', '')
    confirm_update = d.get('confirm_update', False)

    committee = None
    if committee_id:
        committee, error = resolve_grade_committee(
            source, pid, ctype, semester=semester, committee_id=committee_id
        )
        if error:
            return {'ok': False, 'error': error, 'status': status.HTTP_400_BAD_REQUEST}

    if committee is not None:
        allowed_individual_grader = committee.chair_id == user.id
    else:
        allowed_individual_grader = _check_grader_permission(user, source, pid, ctype)

    if not allowed_individual_grader:
        return {'ok': False, 'error': 'أنت لست رئيس اللجنة المسؤولة عن هذا المشروع.',
                'status': status.HTTP_403_FORBIDDEN}

    if committee is None:
        committee, error = resolve_grade_committee(
            source, pid, ctype, semester=semester, committee_id=None
        )
        if error:
            return {'ok': False, 'error': error, 'status': status.HTTP_400_BAD_REQUEST}

    mode = CommitteeGradingMode.objects.filter(committee=committee).first()
    if mode and mode.collective:
        return {
            'ok': False,
            'error': 'التقييم الجماعي مُفعَّل لهذه اللجنة؛ يجب إدخال تقييم كل عضو عبر مسار التقييم الجماعي.',
            'status': status.HTTP_409_CONFLICT,
        }

    committee_id = committee.id
    semester = committee.semester or semester

    requested_student_ids = [item['student_id'] for item in d['grades']]
    if len(requested_student_ids) != len(set(requested_student_ids)):
        return {'ok': False, 'error': 'لا يجوز تكرار الطالب نفسه في الطلب.',
                'status': status.HTTP_400_BAD_REQUEST}
    existing_students = User.objects.filter(pk__in=requested_student_ids, role='student').in_bulk()
    outside_project = set(existing_students).difference(active_project_student_ids(source, pid))
    if outside_project:
        return {'ok': False, 'error': 'يتضمن الطلب طالباً غير نشط في المشروع المحدد.',
                'status': status.HTTP_400_BAD_REQUEST}

    if not confirm_update:
        existing = ProjectGrade.objects.filter(
            project_source=source, project_id=pid, committee_type=ctype,
            student__in=[item['student_id'] for item in d['grades']],
        ).exists()
        if existing:
            return {
                'ok': False,
                'requires_confirmation': True,
                'status': status.HTTP_409_CONFLICT,
            }

    saved = []
    for item in d['grades']:
        try:
            student = User.objects.get(pk=item['student_id'], role='student')
        except User.DoesNotExist:
            continue

        grade, created = ProjectGrade.objects.get_or_create(
            project_source=source, project_id=pid, committee_type=ctype, student=student,
            defaults={'semester': semester, 'committee_id': committee_id, 'entered_by': user},
        )
        old_main, old_report = grade.score_main, grade.score_report

        grade.score_main   = item['score_main']
        grade.score_report = item.get('score_report')
        grade.notes        = item.get('notes', '')
        grade.entered_by   = user
        if not grade.semester:
            grade.semester = semester
        grade.committee_id = committee_id
        grade.save()

        _log_grade_changes(grade, user, ctype, old_main, old_report)
        saved.append(grade)

    return {'ok': True, 'saved': saved}


# ── Project Grades Detail ─────────────────────────────────────────────────────

def get_project_grades(*, user, source, pid):
    pid = int(pid)
    if source not in VALID_GRADE_PROJECT_SOURCES:
        return {'ok': False, 'error': 'مصدر المشروع غير صالح.', 'status': status.HTTP_400_BAD_REQUEST}

    if is_student(user):
        if not student_belongs_to_project(user, source, pid):
            return {'ok': False, 'error': 'ليس لديك صلاحية.', 'status': status.HTTP_403_FORBIDDEN}
        grades = ProjectGrade.objects.filter(
            project_source=source, project_id=pid, student=user
        ).select_related('student')
    elif is_doctor(user):
        from projects.models import IdeaApplication, StudentIdeaProposal
        can_access = is_dean(user) or Committee.objects.filter(
            project_q_filter(source, pid)
        ).filter(Q(chair=user) | Q(members=user)).exists()
        if source == 'IdeaApplication':
            can_access = can_access or IdeaApplication.objects.filter(pk=pid, idea__doctor=user).exists()
        else:
            can_access = can_access or StudentIdeaProposal.objects.filter(
                Q(pk=pid, supervisor=user) | Q(pk=pid, co_supervisors=user)
            ).exists()
        if not can_access:
            return {'ok': False, 'error': 'ليس لديك صلاحية.', 'status': status.HTTP_403_FORBIDDEN}
        grades = ProjectGrade.objects.filter(
            project_source=source, project_id=pid
        ).select_related('student')
    else:
        return {'ok': False, 'error': 'ليس لديك صلاحية.', 'status': status.HTTP_403_FORBIDDEN}

    report = ProjectReport.objects.filter(project_source=source, project_id=pid).first()

    students_map = {}
    for g in grades:
        sid = g.student_id or 0
        if sid not in students_map:
            students_map[sid] = {
                'student_id':       sid,
                'student_name':     (g.student.get_full_name() or g.student.username) if g.student else '—',
                'student_username': g.student.username if g.student else '—',
                'grades':           {},
                'total_score':      0,
            }
        students_map[sid]['grades'][g.committee_type] = ProjectGradeSerializer(g).data
        students_map[sid]['total_score'] += g.total_score

    return {
        'ok': True,
        'project_source':  source,
        'project_id':      pid,
        'students_grades': list(students_map.values()),
        'report':          report,
    }


# ── Doctor — My Committee Grades ──────────────────────────────────────────────

def get_my_committee_grades(*, user, semester=None):
    if not is_doctor(user):
        return {'ok': False, 'error': 'مسموح للدكاترة فقط.', 'status': status.HTTP_403_FORBIDDEN}

    chaired  = Committee.objects.filter(chair=user)
    membered = Committee.objects.filter(members=user)
    committees = (chaired | membered).distinct()
    if semester:
        committees = committees.filter(semester=semester)

    result = []
    for c in committees:
        is_chair   = c.chair_id == user.id
        mode       = CommitteeGradingMode.objects.filter(committee=c).first()
        collective = mode.collective if mode else False

        # Individual mode is chair-only.  Ordinary members (including HoD)
        # participate only when collective grading is enabled.
        if not is_chair and not collective:
            continue

        result.append(_build_committee_grades_entry(c, user, is_chair, collective))

    return {'ok': True, 'committees': result}


def _build_committee_grades_entry(c, user, is_chair, collective):
    projects_data = []
    for p in c.get_all_projects():
        source, p_id = p['source'], p['id']
        is_final = c.committee_type == 'final_discussion'

        grades_qs = ProjectGrade.objects.filter(
            project_source=source, project_id=p_id, committee_type=c.committee_type,
        ).select_related('student')
        grades_by_student = {g.student_id: g for g in grades_qs}

        report = ProjectReport.objects.filter(project_source=source, project_id=p_id).first()

        my_drafts = {}
        if collective:
            for draft in DoctorGradeDraft.objects.filter(
                committee=c, project_source=source, project_id=p_id,
                committee_type=c.committee_type, doctor=user,
            ).select_related('student'):
                my_drafts[draft.student_id] = draft

        students_with_grades = []
        for s in p.get('students', []):
            if s.get('status', 'active') != 'active':
                continue
            s_id, s_grade, s_draft = s.get('id'), grades_by_student.get(s.get('id')), my_drafts.get(s.get('id'))
            students_with_grades.append({
                'student_id':   s_id,
                'student_name': s.get('name', ''),
                'is_leader':    s.get('is_leader', False),
                'grade':        ProjectGradeSerializer(s_grade).data if s_grade else None,
                'my_draft': {
                    'score_main':   s_draft.score_main,
                    'score_report': s_draft.score_report,
                    'notes':        s_draft.notes,
                } if s_draft else None,
            })

        projects_data.append({
            'id': p_id, 'source': source, 'title': p.get('title', ''),
            'students': students_with_grades,
            'max_score_main':   COMMITTEE_MAX_SCORES.get(c.committee_type, 0),
            'max_score_report': 30 if is_final else 0,
            'report_uploaded':  report is not None,
            'report': ProjectReportSerializer(report).data if report else None,
            'all_graded': (
                len(students_with_grades) > 0
                and all(
                    sw['grade'] is not None
                    and sw['grade'].get('score_main') is not None
                    and (
                        not is_final
                        or sw['grade'].get('score_report') is not None
                    )
                    for sw in students_with_grades
                )
            ),
        })

    return {
        'committee_id': c.id,
        'committee_type': c.committee_type,
        'committee_type_ar': COMMITTEE_TYPE_AR.get(c.committee_type, c.committee_type),
        'department': c.department,
        'department_ar': DEPARTMENT_AR.get(c.department, c.department),
        'semester': c.semester,
        'date': c.date.strftime('%Y-%m-%d') if c.date else None,
        'max_score': COMMITTEE_MAX_SCORES.get(c.committee_type, 0),
        'collective_mode': getattr(CommitteeGradingMode.objects.filter(committee=c).first(), 'collective', False),
        'is_chair': is_chair,
        'projects': projects_data,
    }


# ── Student — My Grades ───────────────────────────────────────────────────────

def get_my_grades(*, user):
    if not is_student(user):
        return {'ok': False, 'error': 'مسموح للطلاب فقط.', 'status': status.HTTP_403_FORBIDDEN}

    from projects.models import ProjectParticipation
    parts = ProjectParticipation.objects.filter(student=user, status='active').select_related(
        'idea_application__idea', 'student_proposal'
    )

    result = []
    for part in parts:
        if part.project_source == 'idea_application' and part.idea_application_id:
            source, pid, title = 'IdeaApplication', part.idea_application_id, part.idea_application.idea.title
            department = part.idea_application.idea.department
        elif part.project_source == 'student_proposal' and part.student_proposal_id:
            source, pid, title = 'StudentIdeaProposal', part.student_proposal_id, part.student_proposal.title
            department = part.student_proposal.department
        else:
            continue

        grades = ProjectGrade.objects.filter(project_source=source, project_id=pid, student=user)
        report = ProjectReport.objects.filter(project_source=source, project_id=pid).first()

        grades_by_type = {g.committee_type: ProjectGradeSerializer(g).data for g in grades}
        total = sum(g.total_score for g in grades)

        result.append({
            'project_source': source, 'project_id': pid, 'project_title': title,
            'department': department, 'department_ar': DEPARTMENT_AR.get(department, department),
            'role': part.role, 'grades': grades_by_type, 'total_score': total,
            'max_total': 100, 'report_uploaded': report is not None,
            'report': ProjectReportSerializer(report).data if report else None,
        })

    return {'ok': True, 'projects': result}


# ── Dean — Summary / Excel / Word ─────────────────────────────────────────────

def get_project_info(source, pid):
    from projects.models import IdeaApplication, StudentIdeaProposal, ProjectParticipation
    title = department = project_type = ''
    students = []

    if source == 'IdeaApplication':
        app = IdeaApplication.objects.select_related('idea').filter(pk=pid).first()
        if app:
            title, department, project_type = app.idea.title, app.idea.department, app.idea.project_type
    else:
        prop = StudentIdeaProposal.objects.filter(pk=pid).first()
        if prop:
            title, department, project_type = prop.title, prop.department, prop.project_type

    parts = ProjectParticipation.objects.filter(status='active').select_related('student')
    parts = parts.filter(idea_application_id=pid) if source == 'IdeaApplication' else parts.filter(student_proposal_id=pid)

    for p in parts:
        students.append({
            'name': p.student.get_full_name() or p.student.username,
            'id': p.student.username, 'pk': p.student.pk,
            'is_leader': p.role == 'leader',
        })

    return title, students, department, project_type


def build_grades_summary(semester, department=None, project_type_filter=None, committee_type_filter=None):
    """
    بناء ملخص العلامات: كل مشروع → كل طالب → علاماته في كل لجنة.
    """
    if committee_type_filter and committee_type_filter not in ALL_COMMITTEE_TYPES:
        return {'projects': [], 'count': 0, 'error': 'committee_type غير صالح'}

    active_committee = None
    if committee_type_filter:
        active_committee = {
            'type': committee_type_filter,
            'label': committee_type_filter,
            'max_score': COMMITTEE_MAX_SCORES.get(committee_type_filter),
        }

    grade_qs = ProjectGrade.objects.select_related('student').all()
    if semester:
        grade_qs = grade_qs.filter(semester=semester)
    if committee_type_filter:
        grade_qs = grade_qs.filter(committee_type=committee_type_filter)

    project_student_grades = defaultdict(lambda: defaultdict(dict))
    for g in grade_qs:
        project_student_grades[(g.project_source, g.project_id)][g.student_id or 0][g.committee_type] = g

    rows = []
    for (source, pid), students_data in sorted(project_student_grades.items()):
        title, all_students, proj_department, proj_type = get_project_info(source, pid)

        if department and proj_department != department:
            continue
        if project_type_filter and proj_type != project_type_filter:
            continue

        for s_id, grades_by_type in students_data.items():
            s1, s2 = grades_by_type.get('seminar_1'), grades_by_type.get('seminar_2')
            tc, fd = grades_by_type.get('technical'), grades_by_type.get('final_discussion')

            if committee_type_filter:
                selected = grades_by_type.get(committee_type_filter)
                if not selected:
                    continue
                student_name = (selected.student.get_full_name() or selected.student.username) if selected.student else '—'
                student_uid  = selected.student.username if selected.student else '—'
                rows.append({
                    'project_source': source, 'project_id': pid, 'title': title,
                    'department': proj_department, 'project_type': proj_type,
                    'student_name': student_name, 'student_uid': student_uid,
                    'committee_type': committee_type_filter,
                    'score': selected.score_main if selected.score_main is not None else None,
                })
                continue

            any_grade = s1 or s2 or tc or fd
            student_name = (any_grade.student.get_full_name() or any_grade.student.username) if any_grade and any_grade.student else '—'
            student_uid  = any_grade.student.username if any_grade and any_grade.student else '—'

            total = sum([
                (s1.score_main or 0) if s1 else 0,
                (s2.score_main or 0) if s2 else 0,
                (tc.score_main or 0) if tc else 0,
                (fd.score_main or 0) if fd else 0,
                (fd.score_report or 0) if fd else 0,
            ])

            rows.append({
                'project_source': source, 'project_id': pid, 'title': title,
                'department': proj_department, 'project_type': proj_type,
                'student_name': student_name, 'student_uid': student_uid,
                'seminar_1': s1.score_main if s1 else None,
                'seminar_2': s2.score_main if s2 else None,
                'technical': tc.score_main if tc else None,
                'final_discussion': fd.score_main if fd else None,
                'report': fd.score_report if fd else None,
                'total': total,
            })

    return {'projects': rows, 'count': len(rows), 'active_committee': active_committee}


def get_hod_grades_summary(*, user, semester=None, project_type_filter=None, committee_type_filter=None):
    """رئيس القسم يرى علامات مشاريع قسمه فقط (العميد يرى الكل)."""
    if not is_hod(user):
        return {'ok': False, 'error': 'مسموح لرئيس القسم فقط.', 'status': status.HTTP_403_FORBIDDEN}

    department = getattr(user, 'department', None)
    if not department and not is_dean(user):
        return {'ok': False, 'error': 'حساب رئيس القسم غير مرتبط بقسم.', 'status': status.HTTP_400_BAD_REQUEST}

    dept_filter = None if is_dean(user) else department
    summary = build_grades_summary(semester, dept_filter, project_type_filter, committee_type_filter)
    return {'ok': True, 'summary': summary}


def build_excel_export(semester, department=None, project_type_filter=None, committee_type_filter=None):
    """بناء ملف Excel بكل علامات المشاريع."""
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    except ImportError:
        raise ImportError('openpyxl مطلوب لتصدير Excel')

    summary = build_grades_summary(semester, department, project_type_filter, committee_type_filter)
    rows = summary['projects']

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'علامات المشاريع'
    ws.sheet_view.rightToLeft = True

    header_fill = PatternFill('solid', fgColor='4F46E5')
    alt_fill    = PatternFill('solid', fgColor='F0F0FF')
    thin_border = Border(left=Side(style='thin'), right=Side(style='thin'),
                          top=Side(style='thin'), bottom=Side(style='thin'))

    if committee_type_filter:
        committee_label = COMMITTEE_TYPE_AR.get(committee_type_filter, committee_type_filter)
        max_score = COMMITTEE_MAX_SCORES.get(committee_type_filter, "N/A")
        headers = ['اسم الطالب', 'الرقم الجامعي', 'عنوان المشروع', f'{committee_label} /{max_score}']
        col_widths = [30, 18, 40, 20]
    else:
        headers = [
            'رقم المشروع', 'عنوان المشروع', 'القسم', 'الطالب', 'الرقم الجامعي',
            'سيمينار 1 (10)', 'سيمينار 2 (10)', 'لجنة فنية (20)',
            'مناقشة نهائية (30)', 'تقرير (30)', 'المجموع (100)',
        ]
        col_widths = [12, 32, 16, 22, 14, 14, 14, 16, 18, 12, 14]

    for col_idx, (h, w) in enumerate(zip(headers, col_widths), start=1):
        cell = ws.cell(row=1, column=col_idx, value=h)
        cell.font = Font(bold=True, color='FFFFFF', size=11)
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        cell.border = thin_border
        ws.column_dimensions[ws.cell(row=1, column=col_idx).column_letter].width = w
    ws.row_dimensions[1].height = 30

    for row_idx, proj in enumerate(rows, start=2):
        fill = alt_fill if row_idx % 2 == 0 else PatternFill()
        dept_ar = DEPARTMENT_AR.get(proj['department'], proj['department'])

        if committee_type_filter:
            values = [proj['student_name'], proj['student_uid'], proj['title'],
                      proj['score'] if proj.get('score') is not None else '—']
        else:
            values = [
                f"{proj['project_source'][:3]}-{proj['project_id']}", proj['title'], dept_ar,
                proj['student_name'], proj['student_uid'],
                proj['seminar_1'] if proj['seminar_1'] is not None else '—',
                proj['seminar_2'] if proj['seminar_2'] is not None else '—',
                proj['technical'] if proj['technical'] is not None else '—',
                proj['final_discussion'] if proj['final_discussion'] is not None else '—',
                proj['report'] if proj['report'] is not None else '—',
                proj['total'],
            ]

        for col_idx, val in enumerate(values, start=1):
            cell = ws.cell(row=row_idx, column=col_idx, value=val)
            cell.fill = fill
            cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            cell.border = thin_border
            if committee_type_filter and col_idx == 4:
                cell.font = Font(bold=True)
            elif not committee_type_filter and col_idx == 11:
                cell.font = Font(bold=True)
        ws.row_dimensions[row_idx].height = 20

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.read()


def build_word_export(semester, department, project_type_filter, committee_type_filter=None):
    """بناء ملف Word بتنسيق يشابه النموذج الرسمي."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
    except ImportError:
        raise ImportError('python-docx مطلوب لتصدير Word')

    summary = build_grades_summary(semester, department, project_type_filter, committee_type_filter)
    projects = summary['projects']

    doc = Document()
    doc.default_tab_stops.tabs[0].position = Inches(0.5)

    title = doc.add_paragraph('وثيقة علامات مشاريع القسم', style='Heading 1')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size, run.font.bold = Pt(16), True
        run.font.color.rgb = RGBColor(79, 70, 229)

    dept_text = doc.add_paragraph()
    dept_name = DEPARTMENT_AR.get(department, department) if department else 'جميع الأقسام'
    dept_text.add_run(f'القسم: {dept_name}').font.size = Pt(12)
    dept_text.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if project_type_filter:
        t = doc.add_paragraph()
        t.add_run(f'نوع المشروع: {PROJECT_TYPE_AR.get(project_type_filter, project_type_filter)}').font.size = Pt(12)
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if semester:
        s = doc.add_paragraph()
        s.add_run(f'الفصل: {semester}').font.size = Pt(12)
        s.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=10)
    table.style = 'Light Grid Accent 1'
    table.autofit = False
    table.allow_autofit = False

    headers = [
        'المشروع', 'النوع', 'الطالب', 'الرقم الجامعي',
        'سيمينار 1\n/10', 'سيمينار 2\n/10', 'لجنة فنية\n/20',
        'مناقشة نهائية\n/30', 'تقرير\n/30', 'المجموع\n/100',
    ]
    hdr_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        cell = hdr_cells[i]
        cell.text = header_text
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.bold, run.font.size = True, Pt(11)
            run.font.color.rgb = RGBColor(255, 255, 255)
        _set_cell_background(cell, '4F46E5')

    for proj in projects:
        row_cells = table.add_row().cells
        row_cells[0].text = proj['title'][:50]
        row_cells[1].text = PROJECT_TYPE_AR.get(proj['project_type'], proj['project_type'] or '—')
        row_cells[2].text = proj['student_name']
        row_cells[3].text = str(proj['student_uid'])
        row_cells[4].text = str(proj['seminar_1']) if proj['seminar_1'] is not None else '—'
        row_cells[5].text = str(proj['seminar_2']) if proj['seminar_2'] is not None else '—'
        row_cells[6].text = str(proj['technical']) if proj['technical'] is not None else '—'
        row_cells[7].text = str(proj['final_discussion']) if proj['final_discussion'] is not None else '—'
        row_cells[8].text = str(proj['report']) if proj['report'] is not None else '—'
        row_cells[9].text = str(proj['total'])

        for cell in row_cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(10)
            if len(table.rows) % 2 == 0:
                _set_cell_background(cell, 'F8F7FF')

    col_widths = [1.2, 0.8, 1.2, 1.0, 0.8, 0.8, 1.0, 1.0, 0.8, 0.8]
    for i, width in enumerate(col_widths):
        for row in table.rows:
            row.cells[i].width = Inches(width)

    doc.add_paragraph()
    sig_para = doc.add_paragraph()
    sig_para.add_run('التوقيع: _______________________').font.size = Pt(10)

    for section in doc.sections:
        section_properties = section._sectPr
        bidi_element = OxmlElement('w:bidi')
        section_properties.append(bidi_element)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _set_cell_background(cell, fill_color):
    """تعيين لون خلفية للخلية في جدول Word."""
    try:
        from docx.oxml import parse_xml
        shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(
            'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"', fill_color
        ))
        cell._element.get_or_add_tcPr().append(shading_elm)
    except Exception:
        pass


# ── HoD — Collective Grading Mode ─────────────────────────────────────────────

def hod_department_scope(user):
    """قسم الـ HoD — العميد بلا قيود (None = بدون فلترة)."""
    if is_dean(user):
        return None
    return getattr(user, 'department', None)


def list_grading_modes(*, user):
    if not is_hod(user):
        return {'ok': False, 'error': 'مسموح لرئيس القسم فقط.', 'status': status.HTTP_403_FORBIDDEN}

    dept = hod_department_scope(user)
    if not dept and not is_dean(user):
        return {'ok': False, 'error': 'حساب رئيس القسم غير مرتبط بقسم.', 'status': status.HTTP_400_BAD_REQUEST}

    committees_qs = Committee.objects.all()
    if dept:
        committees_qs = committees_qs.filter(department=dept)

    result = []
    for c in committees_qs:
        mode, _ = CommitteeGradingMode.objects.get_or_create(committee=c)
        result.append({
            'committee_id': c.id,
            'committee_type_ar': COMMITTEE_TYPE_AR.get(c.committee_type, c.committee_type),
            'department_ar': DEPARTMENT_AR.get(c.department, c.department),
            'project_type_ar': PROJECT_TYPE_AR.get(c.project_type, c.project_type),
            'semester': c.semester,
            'collective': mode.collective,
            'updated_at': mode.updated_at.isoformat(),
        })

    return {'ok': True, 'committees': result, 'my_department': dept}


def set_grading_mode(*, user, committee_id, collective):
    if not is_hod(user):
        return {'ok': False, 'error': 'مسموح لرئيس القسم فقط.', 'status': status.HTTP_403_FORBIDDEN}

    if committee_id is None or collective is None:
        return {'ok': False, 'error': 'committee_id و collective مطلوبان.', 'status': status.HTTP_400_BAD_REQUEST}

    try:
        committee = Committee.objects.get(pk=committee_id)
    except Committee.DoesNotExist:
        return {'ok': False, 'error': 'اللجنة غير موجودة.', 'status': status.HTTP_404_NOT_FOUND}

    dept = hod_department_scope(user)
    if dept and committee.department != dept:
        return {'ok': False, 'error': 'لا تملك صلاحية تعديل إعدادات لجنة لا تتبع قسمك.',
                'status': status.HTTP_403_FORBIDDEN}

    mode, _ = CommitteeGradingMode.objects.get_or_create(committee=committee)
    mode.collective = bool(collective)
    mode.set_by = user
    mode.save()

    return {'ok': True, 'committee_id': committee.id, 'collective': mode.collective}


# ── Doctor — Draft Grades (Collective Mode) ───────────────────────────────────

def submit_doctor_drafts(*, user, committee_id, source, pid, ctype, semester, grades_data):
    if not is_doctor(user):
        return {'ok': False, 'error': 'مسموح للدكاترة فقط.', 'status': status.HTTP_403_FORBIDDEN}

    if not (committee_id and source and pid and ctype and grades_data):
        return {
            'ok': False,
            'error': 'committee_id, project_source, project_id, committee_type, grades مطلوبة.',
            'status': status.HTTP_400_BAD_REQUEST,
        }
    if not isinstance(grades_data, list):
        return {'ok': False, 'error': 'grades يجب أن تكون قائمة.', 'status': status.HTTP_400_BAD_REQUEST}

    try:
        committee = Committee.objects.prefetch_related('members').get(pk=committee_id)
    except (Committee.DoesNotExist, TypeError, ValueError):
        return {'ok': False, 'error': 'اللجنة غير موجودة.', 'status': status.HTTP_404_NOT_FOUND}

    mode = CommitteeGradingMode.objects.filter(committee=committee).first()
    if not mode or not mode.collective:
        return {
            'ok': False,
            'error': 'وضع التقييم الجماعي غير مُفعَّل لهذه اللجنة.',
            'status': status.HTTP_400_BAD_REQUEST,
        }

    if not user_is_committee_grader(user, committee):
        return {'ok': False, 'error': 'لست عضواً في هذه اللجنة.', 'status': status.HTTP_403_FORBIDDEN}

    normalised, error = _normalise_grade_request(committee, source, pid, ctype, semester)
    if error:
        return {'ok': False, 'error': error, 'status': status.HTTP_400_BAD_REQUEST}
    pid, semester = normalised

    project_student_ids = active_project_student_ids(source, pid)
    if not project_student_ids:
        return {
            'ok': False,
            'error': 'لا يوجد طلاب نشطون مرتبطون بهذا المشروع.',
            'status': status.HTTP_400_BAD_REQUEST,
        }

    max_main = COMMITTEE_MAX_SCORES.get(ctype)
    if max_main is None:
        return {'ok': False, 'error': 'نوع اللجنة غير صالح.', 'status': status.HTTP_400_BAD_REQUEST}

    is_final = ctype == 'final_discussion'
    validated = []
    seen_students = set()

    for index, item in enumerate(grades_data, start=1):
        if not isinstance(item, dict):
            return {
                'ok': False,
                'error': f'بيانات العلامة في السطر {index} غير صالحة.',
                'status': status.HTTP_400_BAD_REQUEST,
            }

        try:
            student_id = int(item.get('student_id'))
        except (TypeError, ValueError):
            return {
                'ok': False,
                'error': f'رقم الطالب في السطر {index} غير صالح.',
                'status': status.HTTP_400_BAD_REQUEST,
            }

        if student_id in seen_students:
            return {
                'ok': False,
                'error': 'لا يجوز إرسال أكثر من علامة للطالب نفسه في الطلب الواحد.',
                'status': status.HTTP_400_BAD_REQUEST,
            }
        seen_students.add(student_id)

        if student_id not in project_student_ids:
            return {
                'ok': False,
                'error': f'الطالب رقم {student_id} ليس عضواً نشطاً في المشروع المحدد.',
                'status': status.HTTP_400_BAD_REQUEST,
            }

        try:
            score_main = int(item.get('score_main'))
        except (TypeError, ValueError):
            return {
                'ok': False,
                'error': f'العلامة الرئيسية للطالب رقم {student_id} غير صالحة.',
                'status': status.HTTP_400_BAD_REQUEST,
            }
        if score_main < 0 or score_main > max_main:
            return {
                'ok': False,
                'error': f'علامة الطالب رقم {student_id} يجب أن تكون بين 0 و {max_main}.',
                'status': status.HTTP_400_BAD_REQUEST,
            }

        score_report = None
        raw_report = item.get('score_report')
        if is_final and raw_report not in (None, ''):
            try:
                score_report = int(raw_report)
            except (TypeError, ValueError):
                return {
                    'ok': False,
                    'error': f'علامة التقرير للطالب رقم {student_id} يجب أن تكون رقماً.',
                    'status': status.HTTP_400_BAD_REQUEST,
                }
            if score_report < 0 or score_report > 30:
                return {
                    'ok': False,
                    'error': f'علامة التقرير للطالب رقم {student_id} يجب أن تكون بين 0 و 30.',
                    'status': status.HTTP_400_BAD_REQUEST,
                }

        validated.append({
            'student_id': student_id,
            'score_main': score_main,
            'score_report': score_report,
            'notes': str(item.get('notes') or ''),
        })

    students = User.objects.filter(pk__in=seen_students, role='student').in_bulk()
    missing_users = seen_students.difference(students.keys())
    if missing_users:
        return {
            'ok': False,
            'error': f'تعذر العثور على الطلاب: {", ".join(map(str, sorted(missing_users)))}.',
            'status': status.HTTP_400_BAD_REQUEST,
        }

    saved = []
    finalized = []
    pending = []
    with transaction.atomic():
        for item in validated:
            student = students[item['student_id']]
            DoctorGradeDraft.objects.update_or_create(
                committee=committee,
                project_source=source,
                project_id=pid,
                student=student,
                committee_type=ctype,
                doctor=user,
                defaults={
                    'score_main': item['score_main'],
                    'score_report': item['score_report'],
                    'notes': item['notes'],
                },
            )
            saved.append(student.id)

            progress = recalculate_average(
                committee, source, pid, student, ctype, semester, user
            )
            if progress['finalized']:
                finalized.append(student.id)
            else:
                pending.append({
                    'student_id': student.id,
                    'submitted_count': progress['submitted_count'],
                    'required_count': progress['required_count'],
                })

    return {
        'ok': True,
        'saved_students': saved,
        'count': len(saved),
        'finalized_students': finalized,
        'pending_students': pending,
    }


def get_doctor_drafts(*, user, committee_id, source, pid, ctype):
    if not is_doctor(user):
        return {'ok': False, 'error': 'مسموح للدكاترة فقط.', 'status': status.HTTP_403_FORBIDDEN}

    if not (committee_id and source and pid and ctype):
        return {
            'ok': False,
            'error': 'committee_id, project_source, project_id, committee_type مطلوبة.',
            'status': status.HTTP_400_BAD_REQUEST,
        }

    try:
        committee = Committee.objects.prefetch_related('members').get(pk=committee_id)
    except (Committee.DoesNotExist, TypeError, ValueError):
        return {'ok': False, 'error': 'اللجنة غير موجودة.', 'status': status.HTTP_404_NOT_FOUND}

    mode = CommitteeGradingMode.objects.filter(committee=committee).first()
    if not mode or not mode.collective:
        return {
            'ok': False,
            'error': 'وضع التقييم الجماعي غير مُفعَّل لهذه اللجنة.',
            'status': status.HTTP_400_BAD_REQUEST,
        }

    if not user_is_committee_grader(user, committee):
        return {'ok': False, 'error': 'لست عضواً في هذه اللجنة.', 'status': status.HTTP_403_FORBIDDEN}

    normalised, error = _normalise_grade_request(committee, source, pid, ctype, '')
    if error:
        return {'ok': False, 'error': error, 'status': status.HTTP_400_BAD_REQUEST}
    pid, _ = normalised

    grader_ids = committee_grader_ids(committee)
    drafts = DoctorGradeDraft.objects.filter(
        committee=committee,
        project_source=source,
        project_id=pid,
        committee_type=ctype,
        doctor_id__in=grader_ids,
    ).select_related('doctor', 'student')

    data = [{
        'doctor_id': d.doctor_id,
        'doctor_name': d.doctor.get_full_name() or d.doctor.username,
        'student_id': d.student_id,
        'student_name': d.student.get_full_name() or d.student.username,
        'score_main': d.score_main,
        'score_report': d.score_report,
        'notes': d.notes,
        'submitted_at': d.submitted_at.isoformat(),
    } for d in drafts]

    return {
        'ok': True,
        'drafts': data,
        'required_graders_count': len(grader_ids),
    }


def _mark_collective_grade_pending(
    committee, source, pid, student, ctype, submitted_count, required_count, triggered_by
):
    """Hide an old partial collective average while the current committee is incomplete."""
    grade = ProjectGrade.objects.filter(
        project_source=source,
        project_id=pid,
        committee_type=ctype,
        student=student,
        committee=committee,
    ).first()
    if not grade:
        return

    # Only reset grades that were generated by the collective averaging workflow.
    notes = grade.notes or ''
    if not (notes.startswith('متوسط ') or notes.startswith('بانتظار اكتمال التقييم الجماعي')):
        return

    old_main = grade.score_main
    old_report = grade.score_report
    grade.score_main = None
    grade.score_report = None
    grade.entered_by = triggered_by
    grade.notes = f'بانتظار اكتمال التقييم الجماعي ({submitted_count}/{required_count})'
    grade.save(update_fields=['score_main', 'score_report', 'entered_by', 'notes', 'updated_at'])

    if old_main is not None:
        GradeAuditLog.objects.create(
            grade=grade,
            changed_by=triggered_by,
            field_changed='score_main (avg)',
            old_value=str(old_main),
            new_value=None,
        )
    if old_report is not None:
        GradeAuditLog.objects.create(
            grade=grade,
            changed_by=triggered_by,
            field_changed='score_report (avg)',
            old_value=str(old_report),
            new_value=None,
        )


def recalculate_average(committee, source, pid, student, ctype, semester, triggered_by):
    """
    يحسب متوسط التقييم الجماعي بعد اكتمال تقييم جميع أعضاء اللجنة.
    في المناقشة النهائية تُحسب علامة التقرير بصورة مستقلة عن رفع ملف التقرير.
    """
    required_grader_ids = committee_grader_ids(committee)
    if not required_grader_ids:
        return {'finalized': False, 'submitted_count': 0, 'required_count': 0}

    drafts = list(DoctorGradeDraft.objects.filter(
        committee=committee,
        project_source=source,
        project_id=pid,
        student=student,
        committee_type=ctype,
        doctor_id__in=required_grader_ids,
    ))

    drafts_by_doctor = {draft.doctor_id: draft for draft in drafts}
    main_completed_ids = {
        doctor_id for doctor_id, draft in drafts_by_doctor.items()
        if draft.score_main is not None
    }
    required_count = len(required_grader_ids)

    if main_completed_ids != required_grader_ids:
        _mark_collective_grade_pending(
            committee, source, pid, student, ctype,
            len(main_completed_ids), required_count, triggered_by,
        )
        return {
            'finalized': False,
            'submitted_count': len(main_completed_ids),
            'required_count': required_count,
        }

    complete_drafts = [drafts_by_doctor[doctor_id] for doctor_id in required_grader_ids]
    avg_main = round(sum(d.score_main for d in complete_drafts) / required_count)

    is_final = ctype == 'final_discussion'
    report_uploaded = is_final and ProjectReport.objects.filter(
        project_source=source, project_id=pid
    ).exists()
    report_completed_ids = {
        doctor_id for doctor_id, draft in drafts_by_doctor.items()
        if draft.score_report is not None
    } if is_final else set()

    avg_report = None
    report_complete = (not is_final) or report_completed_ids == required_grader_ids
    if is_final and report_complete:
        avg_report = round(
            sum(drafts_by_doctor[doctor_id].score_report for doctor_id in required_grader_ids)
            / required_count
        )

    grade, _ = ProjectGrade.objects.get_or_create(
        project_source=source,
        project_id=pid,
        committee_type=ctype,
        student=student,
        defaults={
            'semester': semester,
            'committee': committee,
            'entered_by': triggered_by,
        },
    )
    old_main, old_report = grade.score_main, grade.score_report

    grade.score_main = avg_main
    grade.score_report = avg_report
    grade.entered_by = triggered_by
    grade.semester = committee.semester or semester or grade.semester
    grade.committee = committee

    if is_final and not report_complete:
        grade.notes = (
            f'متوسط {required_count} تقييمات للمناقشة؛ '
            f'بانتظار اكتمال تقييم التقرير ({len(report_completed_ids)}/{required_count})'
        )
    else:
        grade.notes = f'متوسط {required_count} تقييمات مكتملة'
    grade.save()

    if old_main != avg_main:
        GradeAuditLog.objects.create(
            grade=grade,
            changed_by=triggered_by,
            field_changed='score_main (avg)',
            old_value=str(old_main) if old_main is not None else None,
            new_value=str(avg_main),
        )
    if is_final and old_report != avg_report:
        GradeAuditLog.objects.create(
            grade=grade,
            changed_by=triggered_by,
            field_changed='score_report (avg)',
            old_value=str(old_report) if old_report is not None else None,
            new_value=str(avg_report) if avg_report is not None else None,
        )

    finalized = (not is_final) or report_complete
    pending_count = (
        len(report_completed_ids)
        if is_final and not report_complete
        else required_count
    )
    return {
        'finalized': finalized,
        'submitted_count': pending_count,
        'required_count': required_count,
        'score_main': avg_main,
        'score_report': avg_report,
        'report_uploaded': report_uploaded,
        'report_complete': report_complete,
    }
