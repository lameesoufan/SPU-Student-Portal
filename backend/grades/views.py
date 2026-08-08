"""
Grades App — API Views

Endpoints:
  POST   /api/grades/report/upload/          ← الطالب يرفع التقرير
  GET    /api/grades/report/<source>/<id>/   ← جلب معلومات التقرير
  GET    /api/grades/report/<source>/<id>/download/  ← تحميل التقرير (للجنة)
  POST   /api/grades/enter/                  ← رئيس اللجنة يدخل العلامة
  GET    /api/grades/project/<source>/<id>/  ← علامات مشروع معين
  GET    /api/grades/my-committee-grades/    ← الدكتور يرى مشاريعه وعلاماتها
  GET    /api/grades/export/                 ← العميد يصدر Excel
  GET    /api/grades/summary/                ← ملخص علامات (للعميد)
"""
from __future__ import annotations

import io
import os
from pathlib import Path
from datetime import datetime

from django.http import HttpResponse, FileResponse
from django.utils import timezone
from rest_framework import permissions, status
from rest_framework.parsers import MultiPartParser, FormParser, JSONParser
from rest_framework.response import Response
from rest_framework.views import APIView

from committees.models import Committee, COMMITTEE_TYPE_AR, DEPARTMENT_AR
from .models import (
    ProjectGrade, ProjectReport, GradeAuditLog, COMMITTEE_MAX_SCORES,
    CommitteeGradingMode, DoctorGradeDraft,
)
from .serializers import (
    ProjectGradeSerializer, ProjectReportSerializer, EnterGradeSerializer,
)
from .services import (
    get_doctor_drafts as get_doctor_drafts_service,
    resolve_grade_committee,
    submit_doctor_drafts as submit_doctor_drafts_service,
)


# ── Permission helpers ────────────────────────────────────────────────────────

def _is_student(user): return getattr(user, 'role', None) == 'student'
def _is_doctor(user):  return getattr(user, 'role', None) in ('doctor', 'dean', 'hod')
def _is_dean(user):    return getattr(user, 'role', None) == 'dean'
def _is_hod(user):     return getattr(user, 'role', None) in ('hod', 'dean')


VALID_GRADE_PROJECT_SOURCES = {'IdeaApplication', 'StudentIdeaProposal'}
ALLOWED_REPORT_MIME_TYPES = {
    '.pdf': {'application/pdf'},
    '.doc': {'application/msword'},
    '.docx': {'application/vnd.openxmlformats-officedocument.wordprocessingml.document'},
    '.zip': {'application/zip', 'application/x-zip-compressed'},
    '.rar': {'application/vnd.rar', 'application/x-rar-compressed', 'application/x-rar'},
}


def _safe_uploaded_filename(filename):
    """Strip Windows and POSIX path components from an uploaded filename."""
    raw = str(filename or '').replace('\\', '/')
    safe = raw.rsplit('/', 1)[-1].strip()
    return safe[:255] or 'report'


def _validate_report_upload(file):
    safe_name = _safe_uploaded_filename(getattr(file, 'name', ''))
    ext = os.path.splitext(safe_name)[1].lower()
    if ext not in ALLOWED_REPORT_MIME_TYPES:
        return None, 'نوع الملف غير مسموح.'

    content_type = (getattr(file, 'content_type', '') or '').lower().split(';', 1)[0].strip()
    if content_type not in ALLOWED_REPORT_MIME_TYPES[ext]:
        return None, 'نوع محتوى الملف لا يطابق امتداده.'

    return safe_name, None


def _committee_contains_project(committee, source, pid):
    if source == 'IdeaApplication':
        return committee.applications.filter(pk=pid).exists()
    if source == 'StudentIdeaProposal':
        return committee.proposals.filter(pk=pid).exists()
    return False


def _active_project_student_ids(source, pid):
    from .services import active_project_student_ids
    return active_project_student_ids(source, pid)


def _validate_committee_binding(committee_id, source, pid, committee_type, semester=''):
    if not committee_id:
        return None, None
    try:
        committee = Committee.objects.prefetch_related('members').get(pk=committee_id)
    except (Committee.DoesNotExist, TypeError, ValueError):
        return None, 'اللجنة المحددة غير موجودة.'
    if committee.committee_type != committee_type:
        return None, 'نوع اللجنة لا يطابق اللجنة المحددة.'
    if semester and committee.semester and committee.semester != semester:
        return None, 'الفصل الدراسي لا يطابق فصل اللجنة المحددة.'
    if not _committee_contains_project(committee, source, pid):
        return None, 'المشروع لا يتبع اللجنة المحددة.'
    return committee, None


def _doctor_can_view_project_grades(user, source, pid):
    if _is_dean(user):
        return True

    from django.db.models import Q
    committee_qs = Committee.objects.filter(models_Q_for_project(source, pid)).filter(
        Q(chair=user) | Q(members=user)
    )
    if committee_qs.exists():
        return True

    from projects.models import IdeaApplication, StudentIdeaProposal
    if source == 'IdeaApplication':
        return IdeaApplication.objects.filter(pk=pid, idea__doctor=user).exists()
    if source == 'StudentIdeaProposal':
        return StudentIdeaProposal.objects.filter(
            Q(pk=pid, supervisor=user) | Q(pk=pid, co_supervisors=user)
        ).exists()
    return False


def _get_project(source, pid):
    """جلب كائن المشروع بغض النظر عن المصدر."""
    from projects.models import IdeaApplication, StudentIdeaProposal
    if source == 'IdeaApplication':
        return IdeaApplication.objects.filter(pk=pid).first()
    if source == 'StudentIdeaProposal':
        return StudentIdeaProposal.objects.filter(pk=pid).first()
    return None


def _student_belongs_to_project(user, source, pid):
    """هل الطالب عضو أو قائد في المشروع؟"""
    from projects.models import ProjectParticipation
    qs = ProjectParticipation.objects.filter(student=user, status='active')
    if source == 'IdeaApplication':
        return qs.filter(idea_application_id=pid).exists()
    return qs.filter(student_proposal_id=pid).exists()


def _doctor_is_chair_for(user, source, pid, committee_type):
    """هل الدكتور رئيس لجنة من هذا النوع للمشروع؟"""
    qs = Committee.objects.filter(committee_type=committee_type, chair=user)
    if source == 'IdeaApplication':
        return qs.filter(applications__id=pid).exists()
    return qs.filter(proposals__id=pid).exists()


def _doctor_is_member_for(user, source, pid, committee_type):
    """هل الدكتور/رئيس القسم عضو في لجنة من هذا النوع للمشروع؟"""
    from django.db.models import Q
    qs = Committee.objects.filter(committee_type=committee_type)
    qs = qs.filter(Q(chair=user) | Q(members=user))
    if source == 'IdeaApplication':
        return qs.filter(applications__id=pid).exists()
    return qs.filter(proposals__id=pid).exists()


def _doctor_can_access_report(user, source, pid):
    """هل الدكتور رئيس أو عضو في اللجنة النهائية لهذا المشروع؟"""
    from django.db.models import Q
    qs = Committee.objects.filter(committee_type='final_discussion')
    qs = qs.filter(models_Q_for_project(source, pid))
    return qs.filter(
        Q(chair=user) | Q(members=user)
    ).exists()


def models_Q_for_project(source, pid):
    from django.db.models import Q
    if source == 'IdeaApplication':
        return Q(applications__id=pid)
    return Q(proposals__id=pid)


# ── Report Upload ─────────────────────────────────────────────────────────────

class ReportUploadView(APIView):
    """الطالب يرفع تقرير المشروع."""
    permission_classes = [permissions.IsAuthenticated]
    parser_classes = [MultiPartParser, FormParser]

    def post(self, request):
        user = request.user
        if not _is_student(user):
            return Response({'detail': 'مسموح للطلاب فقط.'}, status=status.HTTP_403_FORBIDDEN)

        source = request.data.get('project_source')
        pid    = request.data.get('project_id')
        sem    = request.data.get('semester', '')
        file   = request.FILES.get('file')

        if not (source and pid and file):
            return Response(
                {'detail': 'project_source, project_id, file مطلوبة.'},
                status=status.HTTP_400_BAD_REQUEST,
            )
        if source not in VALID_GRADE_PROJECT_SOURCES:
            return Response({'detail': 'مصدر المشروع غير صالح.'}, status=status.HTTP_400_BAD_REQUEST)

        try:
            pid = int(pid)
        except ValueError:
            return Response({'detail': 'project_id يجب أن يكون رقماً.'}, status=status.HTTP_400_BAD_REQUEST)

        if not _student_belongs_to_project(user, source, pid):
            return Response({'detail': 'لا تملك صلاحية رفع تقرير لهذا المشروع.'}, status=status.HTTP_403_FORBIDDEN)

        # حجم الملف (10 MB max)
        if file.size > 10 * 1024 * 1024:
            return Response({'detail': 'حجم الملف يتجاوز 10 MB.'}, status=status.HTTP_400_BAD_REQUEST)

        safe_name, upload_error = _validate_report_upload(file)
        if upload_error:
            return Response({'detail': upload_error}, status=status.HTTP_400_BAD_REQUEST)
        file.name = safe_name

        report, created = ProjectReport.objects.get_or_create(
            project_source=source,
            project_id=pid,
            defaults={'semester': sem, 'uploaded_by': user},
        )
        # حذف الملف القديم إن وجد
        if not created and report.file:
            try:
                report.file.delete(save=False)
            except Exception:
                pass

        report.file          = file
        report.original_name = safe_name
        report.file_size     = file.size
        report.semester      = sem or report.semester
        report.uploaded_by   = user
        report.save()

        return Response(ProjectReportSerializer(report, context={'request': request}).data,
                        status=status.HTTP_201_CREATED if created else status.HTTP_200_OK)


# ── Report Info & Download ────────────────────────────────────────────────────

class ReportDetailView(APIView):
    """جلب معلومات التقرير (الطالب أو اللجنة النهائية)."""
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request, source, pid):
        if source not in VALID_GRADE_PROJECT_SOURCES:
            return Response({'detail': 'مصدر المشروع غير صالح.'}, status=status.HTTP_400_BAD_REQUEST)
        user = request.user
        pid  = int(pid)

        # الطالب يرى تقريره فقط
        if _is_student(user):
            if not _student_belongs_to_project(user, source, pid):
                return Response({'detail': 'ليس لديك صلاحية.'}, status=status.HTTP_403_FORBIDDEN)
        elif _is_doctor(user):
            if not (_doctor_can_access_report(user, source, pid) or _is_dean(user)):
                return Response({'detail': 'ليس لديك صلاحية.'}, status=status.HTTP_403_FORBIDDEN)
        else:
            return Response({'detail': 'ليس لديك صلاحية.'}, status=status.HTTP_403_FORBIDDEN)

        try:
            report = ProjectReport.objects.get(project_source=source, project_id=pid)
        except ProjectReport.DoesNotExist:
            return Response({'detail': 'لم يُرفع التقرير بعد.'}, status=status.HTTP_404_NOT_FOUND)

        return Response(ProjectReportSerializer(report, context={'request': request}).data)


class ReportDownloadView(APIView):
    """تحميل ملف التقرير — للجنة النهائية والعميد."""
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request, source, pid):
        if source not in VALID_GRADE_PROJECT_SOURCES:
            return Response({'detail': 'مصدر المشروع غير صالح.'}, status=status.HTTP_400_BAD_REQUEST)
        user = request.user
        pid  = int(pid)

        if _is_student(user):
            if not _student_belongs_to_project(user, source, pid):
                return Response({'detail': 'ليس لديك صلاحية.'}, status=status.HTTP_403_FORBIDDEN)
        elif _is_doctor(user):
            if not (_doctor_can_access_report(user, source, pid) or _is_dean(user)):
                return Response({'detail': 'ليس لديك صلاحية.'}, status=status.HTTP_403_FORBIDDEN)
        else:
            return Response({'detail': 'ليس لديك صلاحية.'}, status=status.HTTP_403_FORBIDDEN)

        try:
            report = ProjectReport.objects.get(project_source=source, project_id=pid)
        except ProjectReport.DoesNotExist:
            return Response({'detail': 'لم يُرفع التقرير بعد.'}, status=status.HTTP_404_NOT_FOUND)

        if not report.file:
            return Response({'detail': 'الملف غير موجود.'}, status=status.HTTP_404_NOT_FOUND)

        response = FileResponse(
            report.file.open('rb'),
            as_attachment=True,
            filename=report.original_name,
        )
        return response


# ── Enter Grade (single student) ─────────────────────────────────────────────

class EnterGradeView(APIView):
    """رئيس اللجنة يدخل علامة طالب واحد."""
    permission_classes = [permissions.IsAuthenticated]
    parser_classes = [JSONParser]

    def post(self, request):
        from .serializers import EnterGradeSerializer
        user = request.user
        if not _is_doctor(user):
            return Response({'detail': 'مسموح للدكاترة فقط.'}, status=status.HTTP_403_FORBIDDEN)

        ser = EnterGradeSerializer(data=request.data)
        ser.is_valid(raise_exception=True)
        d = ser.validated_data

        source       = d['project_source']
        pid          = d['project_id']
        ctype        = d['committee_type']
        student_id   = d['student_id']
        committee_id = d.get('committee_id')
        semester     = d.get('semester', '')

        committee = None
        if committee_id:
            committee, committee_error = resolve_grade_committee(
                source, pid, ctype, semester=semester, committee_id=committee_id
            )
            if committee_error:
                return Response({'detail': committee_error}, status=status.HTTP_400_BAD_REQUEST)

        # Individual/direct grading is chair-only. HoD/Dean privileges do not
        # replace the chair role here; committee members participate through
        # the draft endpoint only when collective mode is enabled.
        allowed_individual_grader = (
            committee.chair_id == user.id
            if committee is not None
            else _doctor_is_chair_for(user, source, pid, ctype)
        )
        if not allowed_individual_grader:
            return Response(
                {'detail': 'التقييم الفردي متاح لرئيس اللجنة فقط.'},
                status=status.HTTP_403_FORBIDDEN,
            )

        if committee is None:
            committee, committee_error = resolve_grade_committee(
                source, pid, ctype, semester=semester, committee_id=None
            )
            if committee_error:
                return Response({'detail': committee_error}, status=status.HTTP_400_BAD_REQUEST)

        mode = CommitteeGradingMode.objects.filter(committee=committee).first()
        if mode and mode.collective:
            return Response(
                {'detail': 'التقييم الجماعي مُفعَّل لهذه اللجنة؛ استخدم مسار تقييم أعضاء اللجنة.'},
                status=status.HTTP_409_CONFLICT,
            )

        committee_id = committee.id
        semester = committee.semester or semester

        from django.contrib.auth import get_user_model
        User = get_user_model()
        try:
            student = User.objects.get(pk=student_id, role='student')
        except User.DoesNotExist:
            return Response({'detail': 'الطالب غير موجود.'}, status=status.HTTP_404_NOT_FOUND)

        if student.id not in _active_project_student_ids(source, pid):
            return Response(
                {'detail': 'الطالب ليس عضواً نشطاً في المشروع المحدد.'},
                status=status.HTTP_400_BAD_REQUEST,
            )

        grade, created = ProjectGrade.objects.get_or_create(
            project_source=source,
            project_id=pid,
            committee_type=ctype,
            student=student,
            defaults={'semester': semester, 'committee_id': committee_id, 'entered_by': user},
        )

        # التحقق من التعديل - إذا كانت العلامة موجودة ولم يؤكد المستخدم
        if not created:
            confirm_update = d.get('confirm_update', False)
            if not confirm_update:
                return Response({
                    'requires_confirmation': True,
                    'message': 'العلامة موجودة بالفعل. هل تريد تغيير العلامة بالتأكيد؟',
                    'existing_grade': ProjectGradeSerializer(grade).data,
                }, status=status.HTTP_409_CONFLICT)

        old_main   = grade.score_main
        old_report = grade.score_report

        grade.score_main   = d['score_main']
        grade.score_report = d.get('score_report')
        grade.notes        = d.get('notes', '')
        grade.entered_by   = user
        if not grade.semester:
            grade.semester = semester
        grade.committee_id = committee_id
        grade.save()

        if old_main != grade.score_main:
            GradeAuditLog.objects.create(grade=grade, changed_by=user,
                field_changed='score_main',
                old_value=str(old_main) if old_main is not None else None,
                new_value=str(grade.score_main))
        if ctype == 'final_discussion' and old_report != grade.score_report:
            GradeAuditLog.objects.create(grade=grade, changed_by=user,
                field_changed='score_report',
                old_value=str(old_report) if old_report is not None else None,
                new_value=str(grade.score_report))

        return Response(ProjectGradeSerializer(grade).data,
                        status=status.HTTP_201_CREATED if created else status.HTTP_200_OK)


# ── Enter Bulk Grades (all students of a project at once) ─────────────────────

class EnterBulkGradesView(APIView):
    """رئيس اللجنة يدخل علامات كل طلاب المشروع دفعة واحدة."""
    permission_classes = [permissions.IsAuthenticated]
    parser_classes = [JSONParser]

    def post(self, request):
        from .serializers import EnterBulkGradesSerializer
        from django.contrib.auth import get_user_model
        User = get_user_model()

        user = request.user
        if not _is_doctor(user):
            return Response({'detail': 'مسموح للدكاترة فقط.'}, status=status.HTTP_403_FORBIDDEN)

        ser = EnterBulkGradesSerializer(data=request.data)
        ser.is_valid(raise_exception=True)
        d = ser.validated_data

        source       = d['project_source']
        pid          = d['project_id']
        ctype        = d['committee_type']
        committee_id = d.get('committee_id')
        semester     = d.get('semester', '')
        confirm_update = d.get('confirm_update', False)

        committee = None
        if committee_id:
            committee, committee_error = resolve_grade_committee(
                source, pid, ctype, semester=semester, committee_id=committee_id
            )
            if committee_error:
                return Response({'detail': committee_error}, status=status.HTTP_400_BAD_REQUEST)

        allowed_individual_grader = (
            committee.chair_id == user.id
            if committee is not None
            else _doctor_is_chair_for(user, source, pid, ctype)
        )
        if not allowed_individual_grader:
            return Response(
                {'detail': 'التقييم الفردي متاح لرئيس اللجنة فقط.'},
                status=status.HTTP_403_FORBIDDEN,
            )

        if committee is None:
            committee, committee_error = resolve_grade_committee(
                source, pid, ctype, semester=semester, committee_id=None
            )
            if committee_error:
                return Response({'detail': committee_error}, status=status.HTTP_400_BAD_REQUEST)

        mode = CommitteeGradingMode.objects.filter(committee=committee).first()
        if mode and mode.collective:
            return Response(
                {'detail': 'التقييم الجماعي مُفعَّل لهذه اللجنة؛ استخدم مسار تقييم أعضاء اللجنة.'},
                status=status.HTTP_409_CONFLICT,
            )

        committee_id = committee.id
        semester = committee.semester or semester

        requested_student_ids = [item['student_id'] for item in d['grades']]
        if len(requested_student_ids) != len(set(requested_student_ids)):
            return Response(
                {'detail': 'لا يجوز تكرار الطالب نفسه في طلب العلامات الجماعي.'},
                status=status.HTTP_400_BAD_REQUEST,
            )

        existing_students = User.objects.filter(
            pk__in=requested_student_ids, role='student'
        ).in_bulk()
        active_student_ids = _active_project_student_ids(source, pid)
        outside_project = set(existing_students).difference(active_student_ids)
        if outside_project:
            return Response(
                {'detail': 'يتضمن الطلب طالباً غير نشط في المشروع المحدد.'},
                status=status.HTTP_400_BAD_REQUEST,
            )

        # التحقق من وجود علامات سابقة - إذا كان هناك أي علامة موجودة ولم يؤكد المستخدم
        if not confirm_update:
            existing_grades = ProjectGrade.objects.filter(
                project_source=source,
                project_id=pid,
                committee_type=ctype,
                student__in=[item['student_id'] for item in d['grades']]
            ).exists()
            
            if existing_grades:
                return Response({
                    'requires_confirmation': True,
                    'message': 'توجد علامات مدخلة سابقاً لأحد الطلاب أو أكثر. هل تريد تغيير العلامات بالتأكيد؟',
                }, status=status.HTTP_409_CONFLICT)

        saved = []
        for item in d['grades']:
            try:
                student = User.objects.get(pk=item['student_id'], role='student')
            except User.DoesNotExist:
                continue

            grade, created = ProjectGrade.objects.get_or_create(
                project_source=source,
                project_id=pid,
                committee_type=ctype,
                student=student,
                defaults={'semester': semester, 'committee_id': committee_id, 'entered_by': user},
            )

            old_main   = grade.score_main
            old_report = grade.score_report

            grade.score_main   = item['score_main']
            grade.score_report = item.get('score_report')
            grade.notes        = item.get('notes', '')
            grade.entered_by   = user
            if not grade.semester:
                grade.semester = semester
            grade.committee_id = committee_id
            grade.save()

            if old_main != grade.score_main:
                GradeAuditLog.objects.create(grade=grade, changed_by=user,
                    field_changed='score_main',
                    old_value=str(old_main) if old_main is not None else None,
                    new_value=str(grade.score_main))
            if ctype == 'final_discussion' and old_report != grade.score_report:
                GradeAuditLog.objects.create(grade=grade, changed_by=user,
                    field_changed='score_report',
                    old_value=str(old_report) if old_report is not None else None,
                    new_value=str(grade.score_report))

            saved.append(ProjectGradeSerializer(grade).data)

        return Response({'saved': saved, 'count': len(saved)})


# ── Project Grades Detail ─────────────────────────────────────────────────────

class ProjectGradesView(APIView):
    """كل علامات مشروع معين مجمّعة per-student."""
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request, source, pid):
        user = request.user
        pid  = int(pid)

        if source not in VALID_GRADE_PROJECT_SOURCES:
            return Response({'detail': 'مصدر المشروع غير صالح.'}, status=status.HTTP_400_BAD_REQUEST)

        if _is_student(user):
            if not _student_belongs_to_project(user, source, pid):
                return Response({'detail': 'ليس لديك صلاحية.'}, status=status.HTTP_403_FORBIDDEN)
            # الطالب يرى علاماته الخاصة فقط
            grades = ProjectGrade.objects.filter(
                project_source=source, project_id=pid, student=user
            ).select_related('student')
        elif _is_doctor(user):
            if not _doctor_can_view_project_grades(user, source, pid):
                return Response({'detail': 'ليس لديك صلاحية.'}, status=status.HTTP_403_FORBIDDEN)
            grades = ProjectGrade.objects.filter(
                project_source=source, project_id=pid
            ).select_related('student')
        else:
            return Response({'detail': 'ليس لديك صلاحية.'}, status=status.HTTP_403_FORBIDDEN)

        report = ProjectReport.objects.filter(project_source=source, project_id=pid).first()

        # تجميع per-student
        students_map = {}
        for g in grades:
            sid = g.student_id or 0
            if sid not in students_map:
                students_map[sid] = {
                    'student_id':       sid,
                    'student_name':     g.student.get_full_name() or g.student.username if g.student else '—',
                    'student_username': g.student.username if g.student else '—',
                    'grades':           {},
                    'total_score':      0,
                }
            students_map[sid]['grades'][g.committee_type] = ProjectGradeSerializer(g).data
            students_map[sid]['total_score'] += g.total_score

        return Response({
            'project_source':  source,
            'project_id':      pid,
            'students_grades': list(students_map.values()),
            'report_uploaded': report is not None,
            'report': ProjectReportSerializer(report, context={'request': request}).data if report else None,
        })


# ── Doctor — My Committee Grades ──────────────────────────────────────────────

class MyCommitteeGradesView(APIView):
    """الدكتور يرى لجانه (كرئيس) مع مشاريعها وعلامات كل طالب."""
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        user = request.user
        if not _is_doctor(user):
            return Response({'detail': 'مسموح للدكاترة فقط.'}, status=status.HTTP_403_FORBIDDEN)

        semester = request.query_params.get('semester')
        # رئيس اللجنة أو عضو فيها (للوضع الجماعي)
        chaired  = Committee.objects.filter(chair=user)
        membered = Committee.objects.filter(members=user)
        committees = (chaired | membered).distinct()
        if semester:
            committees = committees.filter(semester=semester)

        result = []
        for c in committees:
            is_chair     = c.chair_id == user.id
            mode         = CommitteeGradingMode.objects.filter(committee=c).first()
            collective   = mode.collective if mode else False

            # الفردي لرئيس اللجنة فقط؛ الأعضاء يظهر لهم الإدخال فقط عند
            # تفعيل الوضع الجماعي، بغض النظر عن كون العضو Doctor أو HoD.
            if not is_chair and not collective:
                continue

            projects_data = []
            for p in c.get_all_projects():
                source = p['source']
                p_id   = p['id']
                is_final = c.committee_type == 'final_discussion'

                # جلب علامات هذا المشروع في هذه اللجنة مجمّعة per-student
                grades_qs = ProjectGrade.objects.filter(
                    project_source=source,
                    project_id=p_id,
                    committee_type=c.committee_type,
                ).select_related('student')

                # بناء map: student_id → grade
                grades_by_student = {g.student_id: g for g in grades_qs}

                report = ProjectReport.objects.filter(
                    project_source=source, project_id=p_id
                ).first()

                # جلب الـ drafts في الوضع الجماعي (لهذا الطبيب تحديداً)
                my_drafts = {}
                if collective:
                    for draft in DoctorGradeDraft.objects.filter(
                        committee=c,
                        project_source=source,
                        project_id=p_id,
                        committee_type=c.committee_type,
                        doctor=user,
                    ).select_related('student'):
                        my_drafts[draft.student_id] = draft

                # قائمة الطلاب النشطين مع علاماتهم
                students_with_grades = []
                for s in p.get('students', []):
                    if s.get('status', 'active') != 'active':
                        continue
                    s_id    = s.get('id')
                    s_grade = grades_by_student.get(s_id)
                    s_draft = my_drafts.get(s_id)
                    students_with_grades.append({
                        'student_id':       s_id,
                        'student_name':     s.get('name', ''),
                        'is_leader':        s.get('is_leader', False),
                        'grade':            ProjectGradeSerializer(s_grade).data if s_grade else None,
                        'my_draft':         {
                            'score_main':   s_draft.score_main,
                            'score_report': s_draft.score_report,
                            'notes':        s_draft.notes,
                        } if s_draft else None,
                    })

                projects_data.append({
                    'id':              p_id,
                    'source':          source,
                    'title':           p.get('title', ''),
                    'students':        students_with_grades,
                    'max_score_main':  COMMITTEE_MAX_SCORES.get(c.committee_type, 0),
                    'max_score_report': 30 if is_final else 0,
                    'report_uploaded': report is not None,
                    'report': ProjectReportSerializer(report, context={'request': request}).data if report else None,
                    'all_graded': (
                        len(students_with_grades) > 0
                        and all(
                            sw['grade'] is not None
                            and sw['grade'].get('score_main') is not None
                            and (
                                not is_final
                                or sw['grade'].get('score_report') is not None
                            )
                            for sw in students_with_grades
                        )
                    ),
                })

            result.append({
                'committee_id':      c.id,
                'committee_type':    c.committee_type,
                'committee_type_ar': COMMITTEE_TYPE_AR.get(c.committee_type, c.committee_type),
                'department':        c.department,
                'department_ar':     DEPARTMENT_AR.get(c.department, c.department),
                'semester':          c.semester,
                'date':              c.date.strftime('%Y-%m-%d') if c.date else None,
                'max_score':         COMMITTEE_MAX_SCORES.get(c.committee_type, 0),
                'collective_mode':   getattr(
                    CommitteeGradingMode.objects.filter(committee=c).first(),
                    'collective', False
                ),
                'is_chair':          is_chair,
                'projects':          projects_data,
            })

        return Response({'committees': result})


# ── Dean — Summary & Excel Export ────────────────────────────────────────────

class GradesSummaryView(APIView):
    """العميد يرى ملخص علامات كل المشاريع."""
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        if not _is_dean(request.user):
            return Response({'detail': 'مسموح للعميد فقط.'}, status=status.HTTP_403_FORBIDDEN)

        semester = request.query_params.get('semester')
        department = request.query_params.get('department')
        project_type = request.query_params.get('project_type')
        committee_type = request.query_params.get('committee_type')
        return Response(_build_summary(semester, request, department, project_type, committee_type))


class HodGradesSummaryView(APIView):
    """رئيس القسم يرى علامات مشاريع قسمه فقط."""
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        user = request.user
        if not _is_hod(user):
            return Response({'detail': 'مسموح لرئيس القسم فقط.'}, status=status.HTTP_403_FORBIDDEN)

        # جلب قسم رئيس القسم
        department = getattr(user, 'department', None)
        if not department and not _is_dean(user):
            return Response(
                {'detail': 'حساب رئيس القسم غير مرتبط بقسم.'},
                status=status.HTTP_400_BAD_REQUEST,
            )

        semester = request.query_params.get('semester')
        project_type = request.query_params.get('project_type')
        committee_type = request.query_params.get('committee_type')

        # إذا كان عميد يشوف كل الأقسام، وإلا بيشوف قسمه فقط
        dept_filter = None if _is_dean(user) else department

        return Response(_build_summary(semester, request, dept_filter, project_type, committee_type))


class GradesExportView(APIView):
    """العميد يصدّر كل العلامات، ورئيس القسم يصدّر علامات قسمه فقط."""
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        user = request.user
        if not _is_hod(user):
            return Response(
                {'detail': 'مسموح للعميد أو رئيس القسم فقط.'},
                status=status.HTTP_403_FORBIDDEN,
            )

        semester = request.query_params.get('semester')
        project_type = request.query_params.get('project_type')
        committee_type = request.query_params.get('committee_type')
        export_date = request.query_params.get('export_date')

        if _is_dean(user):
            # العميد يستطيع اختيار أي قسم أو تصدير جميع الأقسام.
            department = request.query_params.get('department')
        else:
            # لا نثق بفلتر القسم القادم من الواجهة؛ رئيس القسم مقيد بقسم حسابه.
            department = getattr(user, 'department', None)
            if not department:
                return Response(
                    {'detail': 'حساب رئيس القسم غير مرتبط بقسم.'},
                    status=status.HTTP_400_BAD_REQUEST,
                )

        try:
            content = _build_excel(semester, department, project_type, committee_type, export_date)
        except ValueError as exc:
            return Response({'detail': str(exc)}, status=status.HTTP_400_BAD_REQUEST)

        resp = HttpResponse(
            content,
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        )
        fname = f'grades_{timezone.now():%Y%m%d_%H%M}.xlsx'
        resp['Content-Disposition'] = f'attachment; filename="{fname}"'
        return resp


class HodGradesExportWordView(APIView):
    """رئيس القسم يصدّر علامات مشاريع قسمه بصيغة Word."""
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        if not _is_hod(request.user):
            return Response({'detail': 'مسموح لرئيس القسم فقط.'}, status=status.HTTP_403_FORBIDDEN)

        semester = request.query_params.get('semester')
        project_type = request.query_params.get('project_type')
        committee_type = request.query_params.get('committee_type')
        
        # رئيس القسم يرى فقط قسمه، بينما العميد يستطيع التصدير عبر الأقسام.
        department = getattr(request.user, 'department', None)
        if not department and not _is_dean(request.user):
            return Response(
                {'detail': 'حساب رئيس القسم غير مرتبط بقسم.'},
                status=status.HTTP_400_BAD_REQUEST,
            )

        content = _build_word_grades(semester, department, project_type, committee_type)

        resp = HttpResponse(
            content,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )
        fname = f'hod_grades_{timezone.now():%Y%m%d_%H%M}.docx'
        resp['Content-Disposition'] = f'attachment; filename="{fname}"'
        return resp


# ── Student — My Grades ───────────────────────────────────────────────────────

class MyGradesView(APIView):
    """الطالب يرى علاماته الشخصية لكل لجنة."""
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        user = request.user
        if not _is_student(user):
            return Response({'detail': 'مسموح للطلاب فقط.'}, status=status.HTTP_403_FORBIDDEN)

        from projects.models import ProjectParticipation
        parts = ProjectParticipation.objects.filter(student=user, status='active').select_related(
            'idea_application__idea', 'student_proposal'
        )

        result = []
        for part in parts:
            if part.project_source == 'idea_application' and part.idea_application_id:
                source = 'IdeaApplication'
                pid    = part.idea_application_id
                title  = part.idea_application.idea.title
                department = part.idea_application.idea.department
            elif part.project_source == 'student_proposal' and part.student_proposal_id:
                source = 'StudentIdeaProposal'
                pid    = part.student_proposal_id
                title  = part.student_proposal.title
                department = part.student_proposal.department
            else:
                continue

            # علامات هذا الطالب تحديداً
            grades  = ProjectGrade.objects.filter(
                project_source=source, project_id=pid, student=user
            )
            report  = ProjectReport.objects.filter(project_source=source, project_id=pid).first()

            grades_by_type = {g.committee_type: ProjectGradeSerializer(g).data for g in grades}
            total = sum(g.total_score for g in grades)

            # جلب التشكيلات الحالية مباشرة من جداول اللجان في كل طلب.
            # بهذه الطريقة أي تعديل يجريه العميد على رئيس اللجنة أو أعضائها
            # أو موعدها يظهر تلقائياً للطالب دون تخزين نسخة قديمة في العلامات.
            project_filter = (
                {'applications__id': pid}
                if source == 'IdeaApplication'
                else {'proposals__id': pid}
            )
            committees_qs = (
                Committee.objects
                .filter(**project_filter)
                .select_related('chair', 'room')
                .prefetch_related('members')
                .order_by('committee_type', '-updated_at', '-id')
                .distinct()
            )

            committees_by_type = {}
            for committee in committees_qs:
                # يفترض النظام لجنة واحدة لكل نوع للمشروع. عند وجود بيانات قديمة
                # مكررة نعرض أحدث لجنة عدّلها العميد.
                if committee.committee_type in committees_by_type:
                    continue

                chair = None
                if committee.chair_id:
                    chair = {
                        'id': committee.chair_id,
                        'name': committee.chair.get_full_name() or committee.chair.username,
                    }

                members = [
                    {
                        'id': member.id,
                        'name': member.get_full_name() or member.username,
                    }
                    for member in committee.members.all()
                ]

                committees_by_type[committee.committee_type] = {
                    'id': committee.id,
                    'committee_type': committee.committee_type,
                    'committee_type_ar': COMMITTEE_TYPE_AR.get(
                        committee.committee_type, committee.committee_type
                    ),
                    'department': committee.department,
                    'department_ar': DEPARTMENT_AR.get(committee.department, committee.department),
                    'chair': chair,
                    'members': members,
                    'date': committee.date.isoformat() if committee.date else None,
                    'start_time': (
                        committee.start_time.strftime('%H:%M')
                        if committee.start_time else
                        committee.time.strftime('%H:%M') if committee.time else None
                    ),
                    'end_time': committee.end_time.strftime('%H:%M') if committee.end_time else None,
                    'location': committee.location or '',
                    'room_name': committee.room.name if committee.room_id else None,
                    'status': committee.status,
                    'updated_at': committee.updated_at.isoformat(),
                }

            result.append({
                'project_source':  source,
                'project_id':      pid,
                'project_title':   title,
                'department':      department,
                'department_ar':   DEPARTMENT_AR.get(department, department),
                'role':            part.role,
                'grades':          grades_by_type,
                'committees':      committees_by_type,
                'total_score':     total,
                'max_total':       100,
                'report_uploaded': report is not None,
                'report': ProjectReportSerializer(report, context={'request': request}).data if report else None,
            })

        return Response({'projects': result})


# ── Helper functions ──────────────────────────────────────────────────────────

def _build_summary(semester, request, department=None, project_type_filter=None, committee_type_filter=None):
    """
    بناء ملخص العلامات: كل مشروع → كل طالب → علاماته في كل لجنة.
    يدعم الفلترة حسب القسم ونوع المشروع.
    عند تمرير committee_type_filter يُرجع عمود علامة واحداً (score + committee_type)
    ويُخفي الطلاب الذين ليس لديهم علامة من النوع المختار.
    """
    from committees.models import ALL_COMMITTEE_TYPES
    from .models import COMMITTEE_MAX_SCORES

    # التحقق من صحة فلتر نوع اللجنة
    if committee_type_filter and committee_type_filter not in ALL_COMMITTEE_TYPES:
        return {'projects': [], 'count': 0, 'error': 'committee_type غير صالح'}

    active_committee = None
    if committee_type_filter:
        active_committee = {
            'type': committee_type_filter,
            'label': committee_type_filter,
            'max_score': COMMITTEE_MAX_SCORES.get(committee_type_filter),
        }

    grade_qs = ProjectGrade.objects.select_related('student').all()
    if semester:
        grade_qs = grade_qs.filter(semester=semester)
    if committee_type_filter:
        grade_qs = grade_qs.filter(committee_type=committee_type_filter)

    # تجميع: (source, pid, student_id) → {committee_type: grade}
    from collections import defaultdict
    project_student_grades = defaultdict(lambda: defaultdict(dict))
    # (source, pid) → student_id → committee_type → grade
    for g in grade_qs:
        key     = (g.project_source, g.project_id)
        s_id    = g.student_id or 0
        project_student_grades[key][s_id][g.committee_type] = g

    rows = []
    for (source, pid), students_data in sorted(project_student_grades.items()):
        title, all_students, proj_department, proj_type = _get_project_info(source, pid)

        # الفلترة حسب القسم ونوع المشروع
        if department and proj_department != department:
            continue
        if project_type_filter and proj_type != project_type_filter:
            continue

        for s_id, grades_by_type in students_data.items():
            s1 = grades_by_type.get('seminar_1')
            s2 = grades_by_type.get('seminar_2')
            tc = grades_by_type.get('technical')
            fd = grades_by_type.get('final_discussion')

            # ── وضع فلتر نوع اللجنة: عمود علامة واحد فقط ──
            if committee_type_filter:
                selected = grades_by_type.get(committee_type_filter)
                # إخفاء الطلاب الذين ليس لديهم علامة من النوع المختار
                if not selected:
                    continue
                student_name = '—'
                student_uid  = '—'
                if selected.student:
                    student_name = selected.student.get_full_name() or selected.student.username
                    student_uid  = selected.student.username
                rows.append({
                    'project_source': source,
                    'project_id':     pid,
                    'title':          title,
                    'department':     proj_department,
                    'project_type':   proj_type,
                    'student_name':   student_name,
                    'student_uid':    student_uid,
                    'committee_type': committee_type_filter,
                    'score':          selected.score_main if selected.score_main is not None else None,
                })
                continue

            # اسم الطالب
            any_grade = s1 or s2 or tc or fd
            student_name = '—'
            student_uid  = '—'
            if any_grade and any_grade.student:
                student_name = any_grade.student.get_full_name() or any_grade.student.username
                student_uid  = any_grade.student.username

            total = sum([
                (s1.score_main   or 0) if s1 else 0,
                (s2.score_main   or 0) if s2 else 0,
                (tc.score_main   or 0) if tc else 0,
                (fd.score_main   or 0) if fd else 0,
                (fd.score_report or 0) if fd else 0,
            ])

            rows.append({
                'project_source':    source,
                'project_id':        pid,
                'title':             title,
                'department':        proj_department,
                'project_type':      proj_type,
                'student_name':      student_name,
                'student_uid':       student_uid,
                'seminar_1':         s1.score_main   if s1 else None,
                'seminar_2':         s2.score_main   if s2 else None,
                'technical':         tc.score_main   if tc else None,
                'final_discussion':  fd.score_main   if fd else None,
                'report':            fd.score_report if fd else None,
                'total':             total,
            })

    return {'projects': rows, 'count': len(rows), 'active_committee': active_committee}


def _get_project_info(source, pid):
    from projects.models import IdeaApplication, StudentIdeaProposal, ProjectParticipation
    title = department = project_type = ''
    students = []

    if source == 'IdeaApplication':
        app = IdeaApplication.objects.select_related('idea').filter(pk=pid).first()
        if app:
            title        = app.idea.title
            department   = app.idea.department
            project_type = app.idea.project_type  # semester / graduation_1 / graduation_2
    else:
        prop = StudentIdeaProposal.objects.filter(pk=pid).first()
        if prop:
            title        = prop.title
            department   = prop.department
            project_type = prop.project_type

    parts = ProjectParticipation.objects.filter(
        status='active'
    ).select_related('student')

    if source == 'IdeaApplication':
        parts = parts.filter(idea_application_id=pid)
    else:
        parts = parts.filter(student_proposal_id=pid)

    for p in parts:
        students.append({
            'name':      p.student.get_full_name() or p.student.username,
            'id':        p.student.username,
            'pk':        p.student.pk,
            'is_leader': p.role == 'leader',
        })

    return title, students, department, project_type


def _build_excel(semester, department=None, project_type_filter=None, committee_type_filter=None, export_date=None):
    """إنشاء وثيقة علامات رسمية جاهزة للطباعة وفق الفلاتر المختارة."""
    try:
        import openpyxl
        from openpyxl.drawing.image import Image as XLImage
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
    except ImportError:
        raise ImportError('openpyxl مطلوب لتصدير Excel')

    from committees.models import DEPARTMENT_AR, COMMITTEE_TYPE_AR, PROJECT_TYPE_AR

    # التوافق مع القيمة القديمة القادمة من بعض نسخ الواجهة.
    if project_type_filter == 'semester':
        project_type_filter = 'seasonal'

    if not committee_type_filter:
        raise ValueError('يجب اختيار نوع اللجنة قبل التصدير.')
    if project_type_filter not in ('seasonal', 'graduation_1', 'graduation_2'):
        raise ValueError('يجب اختيار نوع مشروع واحد قبل التصدير.')
    if not export_date:
        raise ValueError('يجب تحديد تاريخ الوثيقة قبل التصدير.')

    try:
        parsed_date = datetime.strptime(export_date, '%Y-%m-%d')
        date_text = parsed_date.strftime('%Y / %m / %d')
    except ValueError:
        raise ValueError('صيغة التاريخ غير صحيحة.')

    summary = _build_summary(semester, None, department, project_type_filter, committee_type_filter)
    rows = summary['projects']

    department_name = DEPARTMENT_AR.get(department, department) if department else 'جميع الأقسام'
    committee_name = COMMITTEE_TYPE_AR.get(committee_type_filter, committee_type_filter)
    project_type_name = PROJECT_TYPE_AR.get(project_type_filter, project_type_filter)
    max_score = COMMITTEE_MAX_SCORES.get(committee_type_filter, '')

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'وثيقة العلامات'
    ws.sheet_view.rightToLeft = True
    ws.sheet_view.showGridLines = False

    # إعداد الصفحة للطباعة على A4.
    ws.page_setup.orientation = 'portrait'
    ws.page_setup.paperSize = ws.PAPERSIZE_A4
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 0
    ws.sheet_properties.pageSetUpPr.fitToPage = True
    ws.print_options.horizontalCentered = True
    ws.page_margins.left = 0.25
    ws.page_margins.right = 0.25
    ws.page_margins.top = 0.35
    ws.page_margins.bottom = 0.45

    dark_blue = '0B2A63'
    gray_fill = PatternFill('solid', fgColor='D9D9D9')
    white_fill = PatternFill('solid', fgColor='FFFFFF')
    thin = Side(style='thin', color='000000')
    medium = Side(style='medium', color='000000')
    table_border = Border(left=thin, right=thin, top=thin, bottom=thin)

    # الأعمدة: اسم الطالب، الرقم الجامعي، عنوان المشروع، العلامة.
    widths = [28, 18, 42, 18]
    for idx, width in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(idx)].width = width

    # شعار الجامعة.
    logo_path = Path(__file__).resolve().parent / 'assets' / 'spu_logo.png'
    if logo_path.exists():
        logo = XLImage(str(logo_path))
        logo.width = 112
        logo.height = 72
        ws.add_image(logo, 'A1')

    # رأس الوثيقة.
    ws.merge_cells('B1:D1')
    ws['B1'] = 'كلية هندسة الذكاء الاصطناعي'
    ws['B1'].font = Font(name='Arial', size=16, bold=True, color=dark_blue)
    ws['B1'].alignment = Alignment(horizontal='center', vertical='center')
    ws.merge_cells('B2:D2')
    ws['B2'] = 'Faculty of Artificial Intelligence Engineering'
    ws['B2'].font = Font(name='Times New Roman', size=14, bold=True, color=dark_blue)
    ws['B2'].alignment = Alignment(horizontal='center', vertical='center')
    ws.row_dimensions[1].height = 32
    ws.row_dimensions[2].height = 25

    ws.merge_cells('A3:B3')
    ws['A3'] = 'رمز الوثيقة: M-09-AI-F6'
    ws.merge_cells('C3:D3')
    ws['C3'] = 'رقم الإصدار: 00'
    ws.merge_cells('A4:B4')
    ws['A4'] = f'تاريخ الإصدار: {date_text}'
    ws.merge_cells('C4:D4')
    ws['C4'] = f'وثيقة علامات عضو لجنة {committee_name}'
    for row in (3, 4):
        for cell in ws[row]:
            cell.font = Font(name='Arial', size=10)
            cell.alignment = Alignment(horizontal='center', vertical='center')
            cell.border = Border(top=thin if row == 3 else Side(style=None), bottom=thin)

    ws.merge_cells('A6:D6')
    ws['A6'] = f'وثيقة علامات عضو لجنة {committee_name}'
    ws['A6'].font = Font(name='Arial', size=17, bold=True)
    ws['A6'].alignment = Alignment(horizontal='center', vertical='center')
    ws.row_dimensions[6].height = 30

    ws.merge_cells('A7:D7')
    ws['A7'] = f'قسم {department_name}'
    ws['A7'].font = Font(name='Arial', size=14, bold=True)
    ws['A7'].alignment = Alignment(horizontal='center', vertical='center')

    # مربعات نوع المشروع، ويظهر رمز صح بجانب الخيار المحدد فقط.
    labels = [
        ('seasonal', 'فصلي'),
        ('graduation_1', 'تخرج 1'),
        ('graduation_2', 'تخرج 2'),
    ]
    checks = ' / '.join(f"{label} {'☑' if key == project_type_filter else '☐'}" for key, label in labels)
    ws.merge_cells('A8:D8')
    ws['A8'] = f'نوع المشروع:  {checks}'
    ws['A8'].font = Font(name='Arial', size=13, bold=True)
    ws['A8'].alignment = Alignment(horizontal='center', vertical='center')
    ws.row_dimensions[8].height = 28

    ws.merge_cells('A9:D9')
    ws['A9'] = f'التاريخ: {date_text}'
    ws['A9'].font = Font(name='Arial', size=11)
    ws['A9'].alignment = Alignment(horizontal='right', vertical='center')

    header_row = 11
    headers = ['اسم الطالب', 'الرقم الجامعي', 'عنوان المشروع', f'العلامة ({max_score} درجات)']
    for col, text in enumerate(headers, 1):
        cell = ws.cell(header_row, col, text)
        cell.font = Font(name='Arial', size=12, bold=True)
        cell.fill = gray_fill
        cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        cell.border = Border(left=medium, right=medium, top=medium, bottom=medium)
    ws.row_dimensions[header_row].height = 42

    first_data_row = header_row + 1
    minimum_rows = 10
    total_rows = max(minimum_rows, len(rows))
    for offset in range(total_rows):
        row_idx = first_data_row + offset
        proj = rows[offset] if offset < len(rows) else None
        values = [
            proj['student_name'] if proj else '',
            proj['student_uid'] if proj else '',
            proj['title'] if proj else '',
            (proj.get('score') if proj and proj.get('score') is not None else '') if proj else '',
        ]
        for col, value in enumerate(values, 1):
            cell = ws.cell(row_idx, col, value)
            cell.font = Font(name='Arial', size=10, bold=(col == 4 and value != ''))
            cell.fill = white_fill
            cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            cell.border = table_border
        ws.row_dimensions[row_idx].height = 27

    signature_row = first_data_row + total_rows + 3
    ws.merge_cells(start_row=signature_row, start_column=1, end_row=signature_row, end_column=2)
    ws.cell(signature_row, 1, 'اسم المدرس وتوقيعه: ................................................')
    ws.cell(signature_row, 1).font = Font(name='Arial', size=12, bold=True)
    ws.cell(signature_row, 1).alignment = Alignment(horizontal='right', vertical='center')

    ws.merge_cells(start_row=signature_row, start_column=3, end_row=signature_row, end_column=4)
    ws.cell(signature_row, 3, f'القسم: {department_name}')
    ws.cell(signature_row, 3).font = Font(name='Arial', size=11)
    ws.cell(signature_row, 3).alignment = Alignment(horizontal='left', vertical='center')
    ws.row_dimensions[signature_row].height = 30

    ws.print_title_rows = f'1:{header_row}'
    ws.print_area = f'A1:D{signature_row + 1}'

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.read()


# ── HoD — Toggle Collective Grading Mode ─────────────────────────────────────

class CommitteeGradingModeView(APIView):
    """
    رئيس القسم يقرأ أو يُغيّر وضع التقييم — لقسمه فقط.

    GET  /api/grades/grading-mode/
         يرجع لجان قسم الـ HoD فقط (department مأخوذ من بيانات المستخدم).

    POST /api/grades/grading-mode/
         { committee_id: int, collective: true|false }
         يُفعّل/يُعطّل الوضع — يُرفض إذا كانت اللجنة لقسم آخر.
    """
    permission_classes = [permissions.IsAuthenticated]
    parser_classes     = [JSONParser]

    def _hod_department(self, user):
        """قسم الـ HoD — العميد يرى الكل (None = no filter)."""
        if _is_dean(user):
            return None          # العميد لا قيود
        return getattr(user, 'department', None)

    def get(self, request):
        user = request.user
        if not _is_hod(user):
            return Response({'detail': 'مسموح لرئيس القسم فقط.'}, status=status.HTTP_403_FORBIDDEN)

        dept = self._hod_department(user)
        if not dept and not _is_dean(user):
            return Response(
                {'detail': 'حساب رئيس القسم غير مرتبط بقسم.'},
                status=status.HTTP_400_BAD_REQUEST,
            )

        committees_qs = Committee.objects.all()
        if dept:
            committees_qs = committees_qs.filter(department=dept)

        from committees.models import COMMITTEE_TYPE_AR, DEPARTMENT_AR, PROJECT_TYPE_AR
        result = []
        for c in committees_qs:
            mode, _ = CommitteeGradingMode.objects.get_or_create(committee=c)
            result.append({
                'committee_id':      c.id,
                'committee_type_ar': COMMITTEE_TYPE_AR.get(c.committee_type, c.committee_type),
                'department_ar':     DEPARTMENT_AR.get(c.department, c.department),
                'project_type_ar':   PROJECT_TYPE_AR.get(c.project_type, c.project_type),
                'semester':          c.semester,
                'collective':        mode.collective,
                'updated_at':        mode.updated_at.isoformat(),
            })

        return Response({'committees': result, 'my_department': dept})

    def post(self, request):
        user = request.user
        if not _is_hod(user):
            return Response({'detail': 'مسموح لرئيس القسم فقط.'}, status=status.HTTP_403_FORBIDDEN)

        committee_id = request.data.get('committee_id')
        collective   = request.data.get('collective')

        if committee_id is None or collective is None:
            return Response(
                {'detail': 'committee_id و collective مطلوبان.'},
                status=status.HTTP_400_BAD_REQUEST,
            )
        if not isinstance(collective, bool):
            return Response(
                {'detail': 'collective يجب أن يكون قيمة منطقية.'},
                status=status.HTTP_400_BAD_REQUEST,
            )

        try:
            committee = Committee.objects.get(pk=committee_id)
        except Committee.DoesNotExist:
            return Response({'detail': 'اللجنة غير موجودة.'}, status=status.HTTP_404_NOT_FOUND)

        # التحقق أن اللجنة تنتمي لقسم الـ HoD
        dept = self._hod_department(user)
        if dept and committee.department != dept:
            return Response(
                {'detail': 'لا تملك صلاحية تعديل إعدادات لجنة لا تتبع قسمك.'},
                status=status.HTTP_403_FORBIDDEN,
            )

        mode, _ = CommitteeGradingMode.objects.get_or_create(committee=committee)
        mode.collective = bool(collective)
        mode.set_by     = user
        mode.save()

        return Response({
            'committee_id': committee.id,
            'collective':   mode.collective,
            'message':      'تم تفعيل التقييم الجماعي.' if mode.collective else 'تم تعطيل التقييم الجماعي.',
        })


# ── Doctor — Submit Draft Grade (Collective Mode) ─────────────────────────────

class DoctorGradeDraftView(APIView):
    """Read and submit per-doctor grades for collective committee grading."""
    permission_classes = [permissions.IsAuthenticated]
    parser_classes = [JSONParser]

    @staticmethod
    def _response(result):
        if result.get('ok'):
            payload = {key: value for key, value in result.items() if key != 'ok'}
            return Response(payload)
        return Response(
            {'detail': result.get('error', 'تعذر تنفيذ الطلب.')},
            status=result.get('status', status.HTTP_400_BAD_REQUEST),
        )

    def post(self, request):
        result = submit_doctor_drafts_service(
            user=request.user,
            committee_id=request.data.get('committee_id'),
            source=request.data.get('project_source'),
            pid=request.data.get('project_id'),
            ctype=request.data.get('committee_type'),
            semester=request.data.get('semester', ''),
            grades_data=request.data.get('grades', []),
        )
        return self._response(result)

    def get(self, request):
        result = get_doctor_drafts_service(
            user=request.user,
            committee_id=request.query_params.get('committee_id'),
            source=request.query_params.get('project_source'),
            pid=request.query_params.get('project_id'),
            ctype=request.query_params.get('committee_type'),
        )
        return self._response(result)


def _build_word_grades(semester, department, project_type_filter, committee_type_filter=None):
    """
    بناء ملف Word بتنسيق يشابه النموذج الرسمي.
    يعرض علامات مشاريع القسم بطريقة احترافية.
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        raise ImportError('python-docx مطلوب لتصدير Word')

    from committees.models import DEPARTMENT_AR, PROJECT_TYPE_AR, COMMITTEE_TYPE_AR
    from projects.models import IdeaApplication, StudentIdeaProposal

    # جلب البيانات
    summary = _build_summary(semester, None, department, project_type_filter, committee_type_filter)
    projects = summary['projects']

    # إنشاء مستند جديد
    doc = Document()
    doc.default_tab_stops.tabs[0].position = Inches(0.5)

    # إضافة العنوان والمعلومات
    title = doc.add_paragraph('وثيقة علامات مشاريع القسم', style='Heading 1')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(79, 70, 229)

    # معلومات القسم
    dept_text = doc.add_paragraph()
    dept_name = DEPARTMENT_AR.get(department, department) if department else 'جميع الأقسام'
    dept_text.add_run(f'القسم: {dept_name}').font.size = Pt(12)
    dept_text.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # نوع المشروع
    if project_type_filter:
        proj_type_text = doc.add_paragraph()
        proj_type_name = PROJECT_TYPE_AR.get(project_type_filter, project_type_filter)
        proj_type_text.add_run(f'نوع المشروع: {proj_type_name}').font.size = Pt(12)
        proj_type_text.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # الفصل الدراسي
    if semester:
        sem_text = doc.add_paragraph()
        sem_text.add_run(f'الفصل: {semester}').font.size = Pt(12)
        sem_text.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # مسافة

    # إنشاء الجدول
    table = doc.add_table(rows=1, cols=10)
    table.style = 'Light Grid Accent 1'
    table.autofit = False
    table.allow_autofit = False

    # رؤوس الأعمدة
    headers = [
        'المشروع', 'النوع', 'الطالب', 'الرقم الجامعي',
        'سيمينار 1\n/10', 'سيمينار 2\n/10', 'لجنة فنية\n/20',
        'مناقشة نهائية\n/30', 'تقرير\n/30', 'المجموع\n/100',
    ]

    hdr_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        cell = hdr_cells[i]
        cell.text = header_text

        # تنسيق رأس العمود
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(255, 255, 255)

        # تلوين الخلية (أزرق)
        _set_cell_background(cell, '4F46E5')

    # إضافة البيانات
    for proj in projects:
        row_cells = table.add_row().cells

        row_cells[0].text = proj['title'][:50]
        row_cells[1].text = PROJECT_TYPE_AR.get(proj['project_type'], proj['project_type'] or '—')
        row_cells[2].text = proj['student_name']
        row_cells[3].text = str(proj['student_uid'])
        row_cells[4].text = str(proj['seminar_1']) if proj['seminar_1'] is not None else '—'
        row_cells[5].text = str(proj['seminar_2']) if proj['seminar_2'] is not None else '—'
        row_cells[6].text = str(proj['technical']) if proj['technical'] is not None else '—'
        row_cells[7].text = str(proj['final_discussion']) if proj['final_discussion'] is not None else '—'
        row_cells[8].text = str(proj['report']) if proj['report'] is not None else '—'
        row_cells[9].text = str(proj['total'])

        # تنسيق البيانات
        for i, cell in enumerate(row_cells):
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(10)

            # تلوين الصفوف بالتناوب
            if len(table.rows) % 2 == 0:
                _set_cell_background(cell, 'F8F7FF')

    # تعيين عرض الأعمدة
    col_widths = [1.2, 0.8, 1.2, 1.0, 0.8, 0.8, 1.0, 1.0, 0.8, 0.8]
    for i, width in enumerate(col_widths):
        for row in table.rows:
            row.cells[i].width = Inches(width)

    # إضافة توقيع
    doc.add_paragraph()
    sig_para = doc.add_paragraph()
    sig_para.add_run('التوقيع: _______________________').font.size = Pt(10)

    # إعداد RTL
    for section in doc.sections:
        section_properties = section._sectPr
        bidi_element = OxmlElement('w:bidi')
        section_properties.append(bidi_element)

    # حفظ في buffer
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _set_cell_background(cell, fill_color):
    """تعيين لون خلفية للخلية في جدول Word."""
    try:
        from docx.oxml import parse_xml
        shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(
            'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"',
            fill_color
        ))
        cell._element.get_or_add_tcPr().append(shading_elm)
    except Exception:
        pass